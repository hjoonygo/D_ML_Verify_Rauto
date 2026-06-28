# -*- coding: utf-8 -*-
# [build_handover_REVoi.py] 세션 핸드오버 6 docx 생성 (캡틴 지시 2026-06-24, §17).
#   공용 Back2TV Guide · REVoi 알파 Guide · Handover · KeyNote · 봇사양서 · Work Order.
import os
from docx import Document
from docx.shared import Pt, RGBColor
DOCS = r"D:\ML\RfRauto\00_Basic_Setup_Package"


def D():
    d = Document(); d.styles["Normal"].font.name = "Malgun Gothic"; d.styles["Normal"].font.size = Pt(10.5); return d
def H(d, t, lv=1): d.add_heading(t, level=lv)
def P(d, t, b=False):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = b; return p
def B(d, items):
    for it in items: d.add_paragraph(it, style="List Bullet")


def save(d, name): path = os.path.join(DOCS, name); d.save(path); print("[저장]", name)


# ───────────────────── 1. 공용 Back2TV 검증 Guide ─────────────────────
d = D()
H(d, "Guide — Back2TV 백테 검증·산출 표준 (공용 방법론) v1", 0)
P(d, "캡틴 지시 2026-06-24. Back2TV = 모든 백테에 쓰는 공용 시스템(특정 봇 전용 아님). 단일출처 = CLAUDE.md §20 + BACKTEST_OUTPUT_SYSTEM.md.", True)
H(d, "0. 정의")
P(d, "Back2TV = 백테 후보 1개 → ①미래참조·환각 100% 원천차단 검증 → ②MDD해제 최고수익 설정(결과+Pine) → ③MDD−25 최고수익 설정(결과+Pine) 을 1회 일괄.")
H(d, "1. ①검증 — 환각·미래참조 원천차단 (verify_*.py)")
P(d, "하나라도 실패하면 그 수치는 폐기. 통과해야 ②③ 진행.")
B(d, [
 "검증엔진만(§15.1) — 봇 로직 재구현 금지.",
 "1m 체결 전수겹침 — 모든 진입·청산가가 그 1분봉 [저,고] 범위 내 + 보유창에서 실제 도달(롱=저≤가, 숏=고≥가).",
 "룩어헤드 규명 — 피봇=확정봉 키(pivots_lr c+right) · 진입가=되돌림 미도달 시 base 폴백(환상 better가 없음) · 신호=봉시작 open·shift(1) 확정 · oi_zscore=롤링z 아핀상쇄(전표본 정규화여도 무해).",
 "★실증: oi_zscore에 롤링z 적용 시 rollz(oi)≡rollz(a·oi+b) 최대차 2.4e-9 → 전표본 정규화 무해 증명.",
])
H(d, "2. ②③ 사이징 — 격리마진·유지증거금·강제청산 실모델")
B(d, [
 "exp=size%/100×lev; mmr(T1 .004/T2 .005, TIER $50k); hsd=1/lev−mmr−SLIP(.0005).",
 "mae≤−hsd → 강제청산 p=−exp×(hsd+COST .0014+|fund|), 아니면 p=R×exp.",
 "★선형 환상 금지: 레버 과하면 청산되어 수익이 안 늚 → 최적 레버가 진짜로 정해짐(실증: 레버30·증거금100%는 단일월 +6040%여도 결국 계좌 $0 소각).",
 "레버×증거금 격자 스윕에서 최고복리 선택 — ②MDD무제한 / ③MDD≤−25.",
])
H(d, "3. Pine 산출 (make_pine.py, //@version=6)")
B(d, [
 "임베드 ≤400거래(초과=최근분만, 제목에 '임베드N/전체M' 표기=침묵금지) · 80거래 슬라이더.",
 "▲롱/▼숏 진입신호(신호봉가) · 수평선=진입평균가 · 선아래 {L/S}평단@수량 · 청산✕ · 수익률(+파랑/−빨강).",
 "가격은 차트 캔들 고저로 클램프(어느 TF든 캔들에 박힘) · 라벨<500. 체결 개별점(✕)은 크기축소로 생략(진입은 삼각형이 겸).",
 "TradingView: BINANCE:BTCUSDT.P · 시간대 UTC · 4h.",
])
H(d, "4. ★경계 (혼동 절대금지)")
P(d, "이 결과는 '전체 36개월 최고 세팅'(과적합 상한·참고용). 환각없음(verify 통과) ≠ 미래보장. '채택'은 held-out·CPCV 표준6(§5.7)·MDD−20 본선(§15·§0) 별도 통과만. Back2TV 책임 = '환각0 검증된 후보 + 눈검증(TV)'까지.", True)
H(d, "5. TIL (이번 세션)")
B(d, [
 "[신뢰95] 격리마진 청산모델이 레버 환상을 제거한다 — 선형모델 +853,337%가 실모델선 MDD−85%/일부세팅 파산으로 정직화.",
 "[신뢰95] full표본 최고수익 ≠ 미래. 같은 봇도 held-out/CPCV로 보면 대폭 낮아짐(예 +1852%→CPCV p25 +12.7%).",
 "[신뢰90] 단일 검증 MDD는 운 좋은 한 구간일 수 있다 — CPCV 표준6가 숨은 MDD(−35%·위반53%)를 드러냄.",
 "도구: back2tv_REVoi.py(탐색+생성)·verify_REVoi.py(검증)·make_pine.py(v6 Pine)·cpcv_reopt.py(36개월 정직 CPCV).",
])
save(d, "Guide_Back2TV_VerificationMethod_v1.docx")

# ───────────────────── 2. REVoi 알파 Guide ─────────────────────
d = D()
H(d, "Guide — REVoi 알파 발견 · 피보 스텝업 v1", 0)
P(d, "캡틴 지시 2026-06-24. REVoi = REV(역추세) + OI(미결제약정) 사용 봇. TS(추세)와 짝.", True)
H(d, "0. REVoi 정의")
P(d, "역추세 신호(mom_24h + OI_z 역방향 합성) → 눌림목(피봇) 정렬 진입 → 피보 스텝업 트레일 청산. 신호는 롤링z·롤링분위(과거만)로 정직.")
H(d, "1. 알파 발견 경위 (수렁↔밧줄)")
B(d, [
 "수렁: 단일 추세(TS)+피보는 held-out OOS≈0(−2%)·2025 음수 = forward 엣지 없음.",
 "밧줄1: 캡틴 '단짠 배합' — 추세(TS)+회귀(REV) 듀얼. 월상관 음수 → REV 단독이 OOS를 캐리(+42%).",
 "밧줄2: ★캡틴 'TS 피보 스텝업을 REV에 붙여라' — REV가 OOS +42%→+76%, 블렌드 OOS +119%로 도약.",
 "결론: REV 단독으로 좁혀 격리마진 청산모델 Back2TV → 36개월 +1852%(MDD−25, full표본 최고)·환각0 검증.",
])
H(d, "2. ★피보 스텝업 범용성 (핵심 TIL)")
P(d, "[신뢰90] 피보 스텝업 청산은 봇을 가리지 않고 알파를 만든다 — 추세봇(TS)뿐 아니라 회귀봇(REV)에 붙여도 대폭 개선. 단 '진입만 되면'이 전제(진입 정렬이 선결).", True)
H(d, "3. 진입 정렬(align)이 관건 TIL")
B(d, [
 "[신뢰80] REV는 '신호 즉시 진입'이 아니라 combo arming → 눌림목 확정 후 진입(align_pivot). 정렬 없이는 붕괴(+12%), 정렬하면 +83%↑.",
 "[신뢰85] REV 신호 정직화: 롤링z·롤링분위(과거 qwin봉)로 full표본 룩어헤드 제거. 이걸로 OOS가 정직해짐(+119%→+29%, 그래도 양수=진짜).",
])
H(d, "4. 성과·한계 (정직)")
B(d, [
 "Back2TV 최고세팅(full표본): 레버3·증거금75% → +1852%/MDD−25%/PF1.36/932거래/강제청산0. 환각0·미래참조0 검증 통과.",
 "★held-out 본선(CPCV 표준6 재최적): p25 +12.7%/yr·80% 폴드 양수 = 수익 엣지 일반화. 그러나 MDD−20 위반 53%·최악−35% = 라이브 부적합.",
 "→ 미결 = MDD 리스크. 해결책 = 레짐인지 리스크컷(Work_Order_REVoi).",
])
save(d, "Guide_AlphaDiscovery_REVoi_FibStepup_v1.docx")

# ───────────────────── 3. Handover ─────────────────────
d = D()
H(d, "Handover — 260624 REVoi · Back2TV", 0)
P(d, "세션: 목표=Regime Detection 이어가기 / 실제성과=REVoi 알파봇 + Back2TV 공용 검증시스템.", True)
H(d, "1. Output of Chat (핵심 산출)")
B(d, [
 "REVoi 알파봇 — 36개월 +1852%(Back2TV 환각0). MDD−20 본선 미통과(미결).",
 "Back2TV 공용 시스템 — 환각 원천차단 검증 + 격리마진 청산모델 + v6 Pine. CLAUDE.md §20·MD 박제.",
 "피보 스텝업을 REV에 적용(bt_full ext_side 확장)으로 알파 도약 입증.",
])
H(d, "2. 산출물 위치")
B(d, [
 "결과: D:\\ML\\RfRauto\\00_WorkHstr\\BackTest_Output\\260624_13_REVoi_MDD25_36mo_v6 · 260624_14_REVoi_MDDfree_36mo_v6 (각 거래원장·월별통합표·v6 Pine·분석txt).",
 "코드: 03_IDEA4Bot\\260623_07_RfRautoAlphaUp\\ (bt_full·back2tv_REVoi·verify_REVoi·make_pine·cpcv_reopt·blend_* + best_blend.json·back2tv_rev_winners.json).",
 "문서: 00_Basic_Setup_Package\\ (이 Handover·KeyNote·Guide 2종·Hstr_Ver_Up_REVoi·Work_Order_REVoi).",
 "지침: CLAUDE.md §20(Back2TV) · BACKTEST_OUTPUT_SYSTEM.md §7.",
])
H(d, "3. PASS / FAIL")
B(d, [
 "PASS: REVoi 환각·미래참조 0(verify_REVoi 전수통과). Back2TV v6 Pine TV정상(34KB·슬라이더). 재현성(같은config=같은수치).",
 "FAIL/미결: REVoi MDD−20 본선(CPCV 폴드 53% 위반·최악−35%). → Work Order 과제1.",
])
H(d, "4. 다음 1수")
P(d, "REVoi에 레짐인지 리스크컷 적용 → CPCV 폴드 MDD≤−20 달성 → held-out 재검증 → 통과 시 §9 확정후보. (Work_Order_REVoi 과제1~6)")
H(d, "5. 신뢰도표")
B(d, [
 "REVoi 수익 엣지 존재(held-out 양수) — 신뢰 80.",
 "REVoi 환각 없음 — 신뢰 95(verify 전수).",
 "REVoi 라이브 적합 — 신뢰 15(MDD 미해결).",
 "Back2TV 공용 시스템 유효 — 신뢰 90.",
])
save(d, "260624_REVoi_Handover.docx")

# ───────────────────── 4. KeyNote ─────────────────────
d = D()
H(d, "KeyNote — 260624 REVoi · Back2TV (기술 상세)", 0)
H(d, "1. bt_full.gen_trades — REV(ext_side) 확장")
B(d, [
 "ext_side(외부신호)·align_pivot·use_trend_flip·arm_bars 추가. ext_side 주면 REV: 외부신호 방향 arming(arm_bars봉 유효) → 눌림목 확정 후 진입.",
 "TS 경로(ext_side=None) 불변. 스톱캡(cand≤close)·1m체결·실펀딩·현실수수료 공유.",
])
H(d, "2. REV 신호 (rev_side, 롤링정직)")
P(d, "mom=open.pct_change(3)·oi=oi_zscore_24h.shift(1) → 각 롤링z(qwin) → combo=−zm·0.048+(−zo)·0.037 → 롤링분위(q,1−q) 임계로 side. 전부 과거만 = 룩어헤드0.")
H(d, "3. verify_REVoi — 환각검증 결과")
B(d, [
 "진입 체결점 2796/2796(100%)·청산 932/932(100%) 1m범위 내. 청산가 보유창 실도달 미도달 0.",
 "oi_zscore 아핀상쇄 최대차 2.4e-9. → 환각·미래참조 0 확정.",
])
H(d, "4. Back2TV 청산모델 사이징 (back2tv_REVoi)")
P(d, "rauto_paper_engine 1:1 격리마진. 레버[3..30]×증거금[10..100%] 격자 스윕. ②MDD무제한·③MDD≤−25 최고복리.")
H(d, "5. make_pine v6")
P(d, "//@version=6 · array.new<float/int> · 임베드≤400 · 80슬라이더 · 캔들 클램프 · 체결점 생략(삼각형이 base겸) · 라벨<500(80×4).")
H(d, "6. 최종 수치")
B(d, [
 "③ MDD−25: 레버3/증거금75% → +1852%/MDD−25%/PF1.36/청산0/단일최고월+44%.",
 "② MDD해제: 레버13/증거금100% → +853,337%/MDD−85%/PF1.09/청산0 (환상 상한·참고).",
 "CPCV 표준6(재최적): p25 +12.7%/yr·음수폴드20%·MDD위반53%·최악−35%.",
])
save(d, "260624_REVoi_KeyNote.docx")

# ───────────────────── 5. 봇 사양서 ─────────────────────
d = D()
H(d, "Hstr_Ver_Up — REVoi 봇 사양서", 0)
P(d, "REVoi = REV(역추세) + OI(미결제약정). 살아있는 사양서. 지우지 말고 대체(ADR 누적).", True)
H(d, "ADR-001 (2026-06-24) — REVoi 신설")
B(d, [
 "문제: 단일 추세봇 held-out OOS≈0. 회귀형 알파 필요.",
 "결정: REVoi 채택 — 역추세 combo(mom+OI) 신호 + 눌림목 정렬 진입 + 피보 스텝업 청산.",
 "근거: 피보 스텝업을 REV에 붙여 OOS +42→+76%, CPCV p25 +12.7%(held-out 양수). 환각0 검증.",
 "상태: 유력 후보(미확정). 미결=MDD−20 위반(CPCV 53%). §9 미등재.",
])
H(d, "사양 (현재)")
B(d, [
 "신호: rev_tf 4h, mom=open.pct_change(3)+oi_zscore_24h, 롤링z/롤링분위(qwin·q).",
 "진입: combo arming(arm봉) → 눌림목(pivot_tf·N) 확정 시 분할진입(피보 되돌림, 미도달=base폴백).",
 "청산: 피보 스텝업 트레일(눌림목 갱신마다 SL 계단상향, 스톱캡). 추세전환 청산 off.",
 "비용: maker2/taker4/스프1bp·실펀딩·1m체결·격리마진 강제청산.",
 "Back2TV 최고세팅(full표본 참고): 레버3·증거금75%·MDD−25 → +1852%.",
])
save(d, "Hstr_Ver_Up_REVoi.docx")

# ───────────────────── 6. Work Order (md→docx) ─────────────────────
d = D()
H(d, "Work Order — REVoi (2026-06-24)", 0)
P(d, "REVoi = REV + OI 봇. 미결 리스크 + 연구과제.", True)
H(d, "맥락")
P(d, "REVoi 36개월 +1852%(Back2TV 환각0). 단 MDD−20 본선 미통과(CPCV 폴드 53% 위반·최악−35%) = 라이브 부적합. 알파는 진짜, 리스크·정밀화가 과제.")
H(d, "과제 1 [미결·즉시] — MDD−20 리스크 해결")
B(d, ["레짐인지 리스크컷(고변동 구간 노출↓)으로 CPCV 폴드 MDD≤−20.",
       "완료조건: CPCV 표준6 전 폴드 MDD≤−20 + p25>0.",
       "재평가: 해결 후 held-out 재검증 통과 시 §9 확정후보 승격."])
H(d, "과제 2 [연구] — TV 기반 진입/청산 패턴 정밀화")
P(d, "TV에서 REVoi 진입·청산 직접 연구 → ①문제점 발굴 ②알파 올릴 패턴 인식(승/패 패턴 규칙화).")
H(d, "과제 3 [연구] — 레짐 적용 (원 세션 목표)")
P(d, "Regime detection 적용 — 추세/횡보/고변동 구분해 진입·사이징·청산 차등(과제1 리스크컷과 연결).")
H(d, "과제 4 [연구] — OI 성격 규명 → 구조레벨 적용")
P(d, "OI(미결제약정) 성격 규명(돌파/저지 작용) → OrderBlock·POC 등 구조레벨에 결합해 돌파/저지 확률↑.")
H(d, "과제 5 [연구] — OB 양면 전략")
P(d, "OB 앞 돌파 실패=그 앞에서 수익실현(저지), 돌파=재진입 짧게 먹기(breakout 스캘핑). OB 양 시나리오 수익화.")
H(d, "과제 6 [연구] — HTF→LTF 스텝업 혼합")
P(d, "HTF(상위TF) 피보 스텝업 이후 LTF(하위TF) 스텝업 혼합. 큰 흐름은 HTF, 정밀 청산/추격은 LTF.")
H(d, "재평가 시점")
P(d, "과제1=다음 세션 즉시. 과제2~6=라이브 후보 확정 후 알파강화 단계, 모두 Back2TV(공용) 검증 후 채택.")
save(d, "Work_Order_REVoi_20260624.docx")

print("\n[완료] 6 docx 생성 →", DOCS)

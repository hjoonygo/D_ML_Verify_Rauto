# -*- coding: utf-8 -*-
# [파일명] build_guide_v4.py — Guide v3 + Part 5(운영·반입 수렁 7건) = v4 생성 (1회용, 캡틴 승인 2026-06-13)
import sys
import docx

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

SRC = r"D:\ml\verify\00WorkHstr\00Basic_Setup_Package\Guide_AlphaDiscovery_Method_v3.docx"
DST = r"D:\ml\verify\00WorkHstr\00Basic_Setup_Package\Guide_AlphaDiscovery_Method_v4.docx"

from docx.shared import Pt

d = docx.Document(SRC)
# v3 문서엔 'Heading 1' 빌트인 스타일이 없음 → 굵게+큰글씨 수동 지정으로 제목 표현
p = d.add_paragraph()
r = p.add_run("Part 5. 운영·반입 수렁 (v4 추가 — 07Prj_Ch4 Stg14자동화+Stg16 OpsGuard, 2026-06-13 캡틴 승인)")
r.bold = True
r.font.size = Pt(14)
d.add_paragraph("알파를 찾는 단계가 아니라, 찾은 알파를 무인으로 돌리는 단계의 수렁들. 형식은 Part 2와 동일: 수렁 → 밧줄 → 잡지말것 → 사례.")

ITEMS = [
("🕳 세션이 어느 머신인지 확인 않고 머신 상태를 바꿈",
 "🪢 밧줄: 작업 전 hostname+핵심경로 Test-Path로 머신 정체부터 증거 확보. 원격 실행 경로(SSH/RDP)가 문서에 없으면 '여기서 그 머신 작업 불가'를 즉시 선언하고 복붙 런북으로 전환.",
 "❌ 잡지말것: 지시문이 AWS라고 가정하고 PC에 schtasks/setx 실행(엉뚱한 머신 오염).",
 "📖 사례: 같은 착오 3턴 반복(PC GramPro16_2510에 AWS 지시) → 매번 증거로 중단, AWS_RDP_RUNBOOK_3MADI.txt(마디별 복붙+정지점) 방식으로 합의."),
("🕳 schtasks SYSTEM 태스크만 침묵(수동 실행은 됨)",
 "🪢 밧줄: ①/TR에 python 풀경로(사용자 설치 python은 SYSTEM PATH에 없음, WindowsApps 스텁 제외) ②환경변수는 전부 setx /M ③진단은 Last Result 코드와 ops 로그 유무의 이분법(로그 안 늘면 미기동, NO-OP 찍히면 env 미가시).",
 "❌ 잡지말것: 태스크 Ready 표시만 보고 가동 중이라 판단.",
 "📖 사례: Telegram_Poll Last Result -2147024894(0x80070002=파일 못 찾음) → 풀경로 재등록+run_daily.bat v7(RAUTO_PY env 우선)로 해소, /status 회신 실측."),
("🕳 원격(캡틴 RDP) 디버깅에서 한 명령씩 왕복하며 시간 소모",
 "🪢 밧줄: 판별표 패턴 — 가설 A/B + 식별용 증거 명령 묶음 + 관찰별 판정·처방 표를 한 번에 전달. 캡틴은 복붙 1회·출력 1회로 끝.",
 "❌ 잡지말것: 증거 없이 가장 그럴듯한 처방부터 적용시키기.",
 "📖 사례: 발신불능 — A(토큰 User영역) vs B(SYSTEM PATH) 판별표 1왕복으로 B 확정."),
("🕳 알림·통신 모듈이 죽으면 본 루틴까지 같이 죽음 / 원인 로그가 없음",
 "🪢 밧줄: ①NO-OP 안전판(자격 미설정=조용히 로그만) ②절대 raise 금지 ③실패 로그에 예외타입+reason 원문 — '다음 실측이 곧 원인 확정'이 되게 설계.",
 "❌ 잡지말것: 발신 실패를 빈 except로 삼키기, 알림 실패로 배치 중단.",
 "📖 사례: alert_telegram v1은 URLError만 찍어 왕복 3회 소모, v2(reason 로그)는 1회 실측으로 해소."),
("🕳 돈·시간에 닿는 설계 트레이드오프를 보고서 단락으로만 흘림",
 "🪢 밧줄: '이 설계면 거래를 다음날 알게 됩니다. 괜찮습니까?'처럼 함의를 질문으로 정면 승격해 승인받기. 캡틴이 새 도구 적응 중일수록 더.",
 "❌ 잡지말것: 한 단락 플래그로 보고했으니 승인된 것으로 간주.",
 "📖 사례: 일일 배치=알림 최대 24h 지연을 플래그만 하고 진행 → 캡틴 격노. 매시간 배치+하트비트로 수습, 실계좌=상주 주문봇 전제 명문화."),
("🕳 pandas가 쓴 CSV를 표준 csv.DictReader로 읽자 첫 컬럼명 깨짐(KeyError)",
 "🪢 밧줄: encoding='utf-8-sig'로 열기(BOM 제거). 산출-소비 코드의 인코딩 짝 확인.",
 "❌ 잡지말것: 컬럼명이 없다고 데이터 파일을 의심부터 하기.",
 "📖 사례: paper_ledger.csv 첫 컬럼 BOM 오염 → S4 샌드박스 테스트가 즉시 검출."),
("🕳 cmd 복붙에서 줄바꿈으로 마지막 옵션이 잘려 인터랙티브 프롬프트 발생",
 "🪢 밧줄: 런북에 '한 줄로' 명시 + 멱등 옵션(/F) 포함 + 프롬프트 발생 시 응답(Y)까지 기재.",
 "❌ 잡지말것: 같은 명령 재전송만 반복.",
 "📖 사례: schtasks /F가 다음 줄로 밀려 Y/N 프롬프트에 /F 입력 → ERROR. Y 한 글자로 해결."),
]
for block in ITEMS:
    for t in block:
        d.add_paragraph(t)
d.save(DST)
print("saved", DST)

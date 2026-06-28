# -*- coding: utf-8 -*-
# [파일명] build_handover_06Prj_Ch8_ControlChart.py — 인수인계보고서+Key노트 docx 생성기 (1회용)
#   채팅: 06Prj_Ch8 Rauto 제어앱(PWA) 차트 정밀화 + 봇 체결시각 보정 + git auto-pull (2026-06-16~17)
import os, sys
import docx
from docx.shared import Pt

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BASE = r"D:\ML\Verify"
PKG = os.path.join(BASE, r"00WorkHstr\00Basic_Setup_Package")
UI = os.path.join(BASE, r"06Prj_Ch8_Plugin_Stg1_TS_Impatient\control_ui")
BOT = os.path.join(BASE, r"06Prj_Ch8_Plugin_Stg1_TS_Impatient\rauto1")
NAME = "06Prj_Ch8_RautoControlPWA_Chart"


def add(doc, style, text):
    if style == "h1":
        doc.add_heading(text, level=1)
    elif style == "h2":
        doc.add_heading(text, level=2)
    elif style == "code":
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(7)
    else:
        doc.add_paragraph(text)


HANDOVER = [
("h1", f"Handover_{NAME} — 인수인계보고서"),
("p", "작성 2026-06-17 | 채팅: 06Prj_Ch8 Rauto 제어앱(PWA) 차트 정밀화 (2026-06-16~17) | 작성자: Claude (Opus 4.8 1M, Claude Code 세션)"),
("p", "★이번 채팅의 한 줄: 새 알파 0건. 폰에 설치한 Rauto 제어앱(PWA)의 '체결 차트'를 트레이딩뷰급으로 정밀화하고, "
      "마커가 캔들과 어긋나던 고질 버그를 봇 데이터 근원에서 + 차트 렌더링에서 2중으로 끝냈다. "
      "그리고 봇·대시보드를 git auto-pull에 편입해 캡틴이 AWS를 만지지 않아도 push만으로 배포되게 만들었다."),

("h1", "★Output of Chat (다음 채팅의 기준점)"),
("p", "① 체결 마커 이탈 버그 종결(전환점): 원인은 봇이 진입/청산 '시각'을 7H봉 라벨로, '가격'은 봉 안쪽 체결값으로 줘서 "
      "시각과 가격이 애초에 안 맞았던 것. 봇(test_Rauto1.py)이 et/xt를 '1분봉에서 그 가격이 실제로 지난 분(分)'으로 역산해 emit하도록 수정. "
      "검증: 진입 11/11, 청산 9~10/11 캔들 정렬(원래 3/11)."),
("p", "② 차트 렌더링도 동시 수정(b20, 전환점): 대시보드가 옛 '가격 재검색(locX)' 방식이라 검색 실패 시 엉뚱한 시간버킷 캔들로 "
      "폴백→십자가가 캔들 밖으로 둥둥 떴다. locX 폐기 → 끝점을 '거래시각에 가장 가까운 캔들'에 배치 + 그 캔들 [저~고]로 "
      "가격 클램프 → 구조적으로 캔들 밖에 못 그림(이탈 0). Y축도 보이는 캔들 범위로만(거래가 확장 제거→캔들 꽉 참). 캡틴 육안 '잘 된다' 확인."),
("p", "③ git auto-pull 운영체제 확립(전환점, RDP 탈출): AWS C:\\RautoRepo에 repo 클론, control_server.py가 180초마다 "
      "git pull --ff-only로 자기 갱신. ★봇도 편입: pull 시 repo의 rauto1/test_Rauto1.py가 바뀌면 C:\\Rauto1로 복사+1회 실행해 "
      "state.json 갱신. → Claude가 push만 하면 대시보드·봇이 AWS 손 안 대고 반영(봇 완전자동은 서버창 1회 재시작으로 활성). 빌드태그 b17→b20 RDP 없이 반영 실증."),
("p", "④ 비용 버전 분리 박제(전환점, ADR-001): §7에 [A]14bp(챔피언/기존TS, Ch7~Ch4_Stg17 재최적까지) · "
      "[B]8bp(성급TS 라이브=지정가진입2+시장청산4+슬립2, 펀딩별도; 경계=Ch4_Stg17 CPCV 202606142318→Ch8_Plugin_Stg1 202606150102) "
      "출처·시각 병기. 라이브 8bp 환산 = +42.75%(4bp +43.96% 대비 -0.84%p). §14 신설(전환점 결과 5곳+6단계 의무기재)."),
("p", "⑤ 확정 알파(§9) 변동 없음. 듀얼 = k0.77 + SW ER>=0.40×0.5 댐핑 → +1097.2% / MDD -16.24% / 표준6 CPCV 위반 0. "
      "이번 채팅은 이 알파를 '보여주고 운영하는' 작업이었음. 06-19 공식 1주 종료 → 기존 vs 성급 §5 직접비교가 다음 채팅 1순위."),

("h1", "0. 파일명·명명"),
("p", f"본 보고서: Handover_{NAME}.docx. 동봉 KeyNote_{NAME}.docx에 코드 전체. 전 파일 영문명(zip 에러 방지). "
      "작업 폴더: D:\\ML\\Verify\\06Prj_Ch8_Plugin_Stg1_TS_Impatient\\ (control_ui\\ = 제어앱, rauto1\\ = 봇 git판)."),

("h1", "1. 작업 히스토리"),
("h2", "1.1 작업 파일 리스트"),
("p", "[control_ui\\ — 제어앱] control_dashboard.html(b20: 자체 캔들차트+TF 15m/1H/4H/1D+EMA+십자선 시:분UTC+ㄱ자 체결마커 near()+클램프) "
      "/ control_server.py(슬롯합치기 aggregate + /cmd 플래그 + ★git auto-pull + ★봇 동기화+재실행) / sw.js(rauto-v20) / manifest.json(fullscreen) "
      "/ start_control_gitpull.bat(클론서 실행+PATH에 Git 추가+RAUTO_GIT_PULL=1)"),
("p", "[rauto1\\ — 봇 git판] test_Rauto1.py(★et/xt를 1분봉 역산으로 보정 + px/trd/equity/wk emit). C:\\Rauto1의 실가동 봇과 동일 — "
      "auto-pull로 동기화됨."),
("p", "[C:\\Rauto1 — AWS/PC 실가동] test_Rauto1.py(동일) / alert_check.py(★실체결가 entry_px/exit_px로 알림 수정, 7H봉종가 표시버그 해소)"),
("p", "[지식·규칙] CLAUDE.md §7(비용 버전병기)·§14(전환점 5곳 의무기재+ADR-001) / LogicCatalog_ByDomain.md v0.2(12분야 로직 메뉴판) / "
      "INDEX 약 8줄 추가(202606160135~202606170300대)"),
("h2", "1.2 코드 목적과 신뢰도 검증 방법"),
("p", "목적: 봇 본체(§8 해시)·신호엔진 무수정 원칙 아래, 폰에서 모든 슬롯을 보고·제어하고, 체결을 트레이딩뷰처럼 차트로 확인. "
      "검증: ①봇 재실행 후 state.json을 파이썬으로 재파싱해 진입/청산이 각 TF 캔들 [저~고]에 드는지 카운트(11/11, 9~10/11) "
      "②대시보드 렌더링 로직(aggPx+near+clmp)을 파이썬으로 그대로 재현해 '이탈 0' 구조 확인 "
      "③캡틴 폰 육안 확인(4 TF 스크린샷) = 최종 PASS 기준."),
("h2", "1.3 결과물 사용법 (캡틴 일상 운용)"),
("p", "폰에서 PWA 열기→체결차트 탭에서 TF(15m/1H/4H/1D) 눌러 캔들·EMA·체결마커 확인. 마커: 가로 가는실선=진입가 레벨, "
      "세로 흰 점선=청산까지 낙차, 양끝 흰 십자=진입/청산 체결점, 수익률 글자(+파랑/-빨강). 숏=분홍, 롱=청색. "
      "긴급상황 버튼→2단 선택창(①정상종료후 ②즉시청산후 자격취소, 실계좌만 중지·가상계속). 새 코드 배포는 Claude가 push만 하면 됨(자동)."),
("h2", "1.4 비망록 — 캡틴이 강조·약속한 것"),
("p", "①작업이 끝날 때마다 반드시 검증하고 보고(메모리 박제) ②AWS는 가능한 한 안 만지게(git auto-pull로 push만) — 새벽 RDP 작업에 매우 지침 "
      "③전환점 결과는 §14대로 5곳 동시기재(CLAUDE.md·Hstr·KeyNote·Guide·INDEX)+6단계 양식 ④비용은 [A]14bp/[B]8bp 정확히 버전 구분(혼동=신뢰 훼손) "
      "⑤'한번에 제대로' — 근원(봇 데이터)부터 고칠 것, 화면 땜질 반복 금지."),

("h1", "2. 시스템 아키텍처"),
("h2", "2.1 전체 그림 (데이터·배포 흐름)"),
("p", "[데이터] Dauto_Collector(1분 수집)→C:\\BinanceData CSV→test_Rauto1.py(매시간 재생, 멱등)→C:\\Rauto1\\state.json(px 15m·trd 체결·equity·wk 통계). "
      "[표시] control_server.py가 C:\\Rauto*\\state.json 합쳐 /state.json 제공→폰 PWA가 fetch→차트·통계 렌더(전부 클라이언트 연산, AWS·봇 타이밍 무영향). "
      "[배포] Claude가 D:\\ML\\Verify에서 git push→AWS control_server가 180초마다 pull→대시보드 즉시·봇은 동기화+재실행으로 반영."),
("h2", "2.2 차트 렌더링 핵심 로직 (코드 전체는 KeyNote)"),
("p", "aggPx(px,m): 15m OHLC를 선택 TF(m분)로 재집계. vis=마지막 110봉. py(v)=가격→y(보이는 캔들 lo~hi 기준), pxi(i)=캔들인덱스→x. "
      "★마커: near(t)=거래시각 t에 가장 가까운 캔들 인덱스, clmp(p,c)=가격을 그 캔들 [저~고]로 강제. "
      "Ye=py(clmp(ep,진입캔들)), Yx=py(clmp(xp,청산캔들)) → 십자가가 절대 캔들 밖으로 안 나감. "
      "ㄱ자: 진입캔들x~청산캔들x 가로선@Ye + 청산캔들x 세로선 Ye→Yx."),
("p", "★봇 _fillms(bar_t, price, win_start): win_start~(bar_t+8h) 1분봉에서 price가 [저~고]를 지난 분 중 bar_t에 가장 가까운 시각을 ms로 반환. "
      "진입은 win_start=진입시각, 청산은 win_start=진입시각(보유기간 전체 — 스톱가는 보유 중 더 낮/높았던 시점값이라 윈도를 넓혀야 찾힘)."),
("h2", "2.3 git auto-pull 운영 (control_server.py)"),
("p", "_git_pull_loop(): 180초마다 git -C C:\\RautoRepo pull --ff-only. 이후 BOT_SRC(repo의 rauto1/test_Rauto1.py)가 BOT_DST(C:\\Rauto1)와 "
      "다르면 복사+python 1회 실행→state.json 갱신. RAUTO_GIT_PULL=1 + RAUTO_REPO 설정 시 기동. "
      "★주의(겪은 함정): 새 control_server.py를 적용하려면 클론이 그 커밋을 pull한 뒤 서버창을 1회 재시작해야 함(이미 떠 있는 서버는 옛 코드 유지). "
      "지금은 클론이 최신이라 다음 1회 재시작이면 봇 자동동기화까지 완전 활성."),

("h1", "3. 인수인계 시 주의점 (점프·미검증 명시)"),
("p", "[구조적 한계 — 확정] 청산 마커 1~2/11은 '모델 스톱레벨'이라 1분봉에 실제로 안 찍힌 가격 → near()가 시간최근 캔들에 놓고 clmp로 캔들 끝에 붙임"
      "(이탈은 없지만 가격이 캔들 끝으로 살짝 보정). 신호엔진(§8 무수정)이 fill을 모델값으로 기록하는 백테 특성이라 봇/차트로 더 못 줄임. 버그 아님."),
("p", "[검증 필요 1 — 신뢰도 55] 봇 자동동기화 완전활성: control_server를 1회 더 재시작해야 새 코드(봇 sync 포함)가 떠서 '봇도 push만으로 자동'이 성립. "
      "이번엔 캡틴이 수동 1줄(copy+run)로 봇 적용함. 다음 채팅에서 서버 재시작 후 'push→봇 자동반영'을 1회 실증할 것."),
("p", "[검증 필요 2 — 신뢰도 90] 15m TF는 보이는 창이 마지막 110봉(약 27.5h)이라 오래된 거래가 안 보일 수 있음. 캡틴 요청 '거래시각까지 x축 패닝'은 "
      "차트 Step3(미구현)에서 해결 예정 — 패닝/핀치/Ctrl+휠 시간줌은 부분 구현, 마커 정렬과 합치는 작업 잔여."),
("p", "[데이터 윈도 주의] 로컬 검증은 11거래(옛 데이터), AWS는 17거래(06-17까지)로 건수가 다름 — 로직 동일하므로 결론 같음. 수치 인용 시 출처(로컬/AWS) 구분할 것."),
("p", "★알파 저장 규칙 리마인드: 알파 확정 시 G:\\...\\06 ChampBot\\00ALPHA_Confirm_Bot에 저장. 이번 채팅 새 알파 0건 — 해당 없음."),

("h1", "4. 개발 토론·문제해결 과정 (상세·코드 전체는 KeyNote)"),
("p", "①마커 이탈 장기전: dot→삼각형→ㄱ자→흰십자선→흰점선 수직→수익률 색구분(v11~v16)으로 화면만 반복 땜질했으나 미해결. "
      "근본원인 규명(봇이 시각=7H봉라벨/가격=봉내체결값으로 불일치) 후 봇 데이터부터 수정→차트 near()+클램프로 2중 종결(b20). 교훈=근원부터."),
("p", "②findCi(전역 가격검색) 치명버그: 비슷한 가격 거래들이 같은 x로 몰리고 진입/청산 역전(+3.5/+3.1/+2.6 같은 시각). locX(시간창 한정)로 1차 수습→"
      "최종 near()(시간최근)로 폐기. ③알림 가격버그: 숏 청산가가 진입가 위로 표시(불가능)→ledger 실체결가(entry_px/exit_px)로 alert_check 수정."),
("p", "④비용 혼동 사고: 14bp(챔피언)와 8bp(성급TS 라이브)를 반복 혼동→캡틴 신뢰 훼손. §7 버전병기+§14 ADR-001로 박제(시간순 INDEX만으론 부족, 단일출처에 버전 병기해야 다음 세션이 정답을 읽음). "
      "⑤git 미설치/PATH 함정: AWS에 git 설치 후 새 cmd창서 PATH 미반영→start_control_gitpull.bat에 'set PATH+=C:\\Program Files\\Git\\cmd' 영구 편입."),

("h1", "5. 다음 채팅 작업 계획 (Output of Chat 기준)"),
("p", "1순위: 06-19 공식 1주 종료 → §5 ★기존 vs 성급 직접비교(장세별·MDD·거래수·휩쓸림·PF·복리수익금). 채택은 CPCV/워크포워드 통과만(§5.6). "
      "2순위: 차트 Step3 — 시간축 패닝(거래시각까지)+핀치/Ctrl+휠 시간줌+자동 Y스케일과 마커 정렬 통합. "
      "3순위: 봇 자동동기화 완전활성 실증(서버 1회 재시작 후 push→봇 자동반영). "
      "4순위(요청): 차트 지표 레지스트리(필요할 때 JS 지표 만들어 등록·on/off) — MACD MTF 포팅은 쉬움(서브페인), 청산 히트맵은 무거움(~400줄). Pine은 우리 차트서 실행 불가(JS 포팅만)."),

("h1", "6. 다음 채팅 시작 시 비판적 검토 사항"),
("p", "①봇 자동동기화: control_server 새 코드가 실제로 떠 있는지(서버 재시작 여부) 먼저 확인 — 안 떠 있으면 push해도 봇 반영 안 됨(주의점 1). "
      "②15m 마커 가시성은 패닝(Step3) 전까지 한계 — 캡틴이 '15m에 거래 안 보임' 지적 시 버그 아니라 윈도 한계임을 먼저 설명. "
      "③06-19 비교는 §5.7 잣대(표준6 CPCV 본선) 엄격 적용 — 관대한 PASS 금지. "
      "④제어앱은 git에 들어가지만 state.json·*.html은 .gitignore라 html은 force-add 필요(이미 추적 중). ⑤AWS CPU/디스크 여유(캡틴 言 느림·용량적음) — auto-pull 180초 주기·봇 재실행 부하 모니터."),

("h1", "부록 A. Key 파일 리스트"),
("p", "[제어앱] control_ui\\control_dashboard.html(b20)·control_server.py·sw.js(v20)·manifest.json·start_control_gitpull.bat "
      "[봇] rauto1\\test_Rauto1.py(=C:\\Rauto1 동일) [알림] C:\\Rauto1\\alert_check.py "
      "[규칙] CLAUDE.md §7·§14, LogicCatalog_ByDomain.md [코드 전체] KeyNote_" + NAME + ".docx"),

("h1", "부록 B. 신뢰도 점수표 (출처 명시)"),
("p", "95: 봇 체결시각 보정 후 진입 11/11 캔들정렬(파이썬 재파싱 실측) / near()+클램프 = 이탈 0(구조적 보장, 렌더링 로직 재현) / "
      "캡틴 폰 육안 '잘 된다' 확인(4 TF) / git auto-pull b17→b20 RDP 없이 반영 실증 / 라이브 8bp = +42.75%(기록 계산)"),
("p", "90: 차트 Step3(패닝/줌) 부분구현 — 마커 정렬과 통합 잔여 / MACD MTF 포팅 난이도 낮음(서브페인 오실레이터)"),
("p", "55: 봇 자동동기화 완전활성(서버 1회 재시작 필요, 다음 채팅 실증) / 청산 마커 1~2건 클램프(모델 스톱레벨, 데이터 한계·버그 아님)"),
]


def build_handover():
    doc = docx.Document()
    for style, text in HANDOVER:
        add(doc, style, text)
    out = os.path.join(PKG, f"Handover_{NAME}.docx")
    doc.save(out)
    return out


KEY_FILES = [
    (UI, "control_dashboard.html"), (UI, "control_server.py"), (UI, "sw.js"),
    (UI, "manifest.json"), (UI, "start_control_gitpull.bat"),
    (BOT, "test_Rauto1.py"),
]


def build_keynote():
    doc = docx.Document()
    add(doc, "h1", f"KeyNote_{NAME} — 핵심·검증 로직 + 코드 전체")
    add(doc, "p", "작성 2026-06-17. 각 파일 머리 주석에 목적·버전·변경사유가 있으니 코드와 함께 읽을 것.")
    add(doc, "h2", "고딩 설명 — 이번 작업의 뼈대")
    add(doc, "p", "봇은 '언제 얼마에 사고팔았는지'를 state.json에 적는다. 문제는 봇이 '시각'을 7시간봉의 시작 라벨로 적고 "
                  "'가격'은 그 봉 안에서 실제 체결된 값으로 적어서, 둘이 어긋났던 것. 그래서 차트에 점을 찍으면 캔들과 동떨어졌다. "
                  "해결: ① 봇이 '그 가격이 실제로 지나간 1분(分)'을 되짚어 시각을 바로잡는다. "
                  "② 차트는 그 시각에 가장 가까운 캔들에 점을 찍고, 가격을 그 캔들의 고가~저가 안으로 '강제(클램프)'한다 — "
                  "그러면 점이 캔들 밖으로 나갈 수가 없다. 두 군데를 같이 고쳐서 끝냈다.")
    for d, fn in KEY_FILES:
        p = os.path.join(d, fn)
        add(doc, "h2", f"[코드 전체] {fn}")
        try:
            body = open(p, encoding="utf-8").read()
        except OSError as e:
            body = f"(읽기 실패: {e})"
        add(doc, "code", body)
    out = os.path.join(PKG, f"KeyNote_{NAME}.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    print("handover:", build_handover())
    print("keynote:", build_keynote())

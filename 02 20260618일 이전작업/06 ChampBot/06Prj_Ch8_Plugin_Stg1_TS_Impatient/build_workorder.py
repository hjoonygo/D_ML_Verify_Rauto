# -*- coding: utf-8 -*-
# Work Order — Rauto 프로덕션(8슬롯·2단 비상대응·텔레그램 제어). §13 미래과제 라우팅.
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"D:\ML\Verify\00WorkHstr\00Basic_Setup_Package\Work_Order_RautoProduction_20260615.docx"
d = docx.Document()


def H(t, sz=15, c=(31, 56, 104)):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = RGBColor(*c); return p


def L(t, b=False):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = b; r.font.size = Pt(10.5); return p


def BULL(t):
    p = d.add_paragraph(style=None); p.paragraph_format.left_indent = Pt(14)
    r = p.add_run("· " + t); r.font.size = Pt(10.5); return p


title = d.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Work Order — Rauto 프로덕션 전환 (8슬롯·2단 비상대응·텔레그램 제어)")
r.bold = True; r.font.size = Pt(17); r.font.color.rgb = RGBColor(31, 56, 104)
sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("06Prj_Ch8 · 2026-06-15 · 작성=Claude Code · 미래과제(라이브 전환 설계 확정)")
rs.font.size = Pt(10); rs.font.color.rgb = RGBColor(110, 110, 110)

H("1. 맥락 (왜)")
L("캡틴 요구(2026-06-15): 실시간 매매봇 상태를 8슬롯 모두 모니터링 + 이머전시 발생 시 1차 자체대응"
  "(진입 급변·슬리피지 → 강제청산 브레이크) + 알람·즉시개입 + 안드로이드(스마트폰) 제어.")
L("06-19 공식 페이퍼 1주 종료 = 채택 결정 시점이지 실돈 런칭 아님. 본 Work Order = '실제 Rauto'의 "
  "프로덕션 아키텍처를 확정하고 테스트넷 착수까지의 설계 단일정의. 기존 Checklist_LiveTransition_Draft(A키/B주문/C가드)를 구체화·계승.")

H("2. 과제 — 목표 아키텍처 (계층)")
L("제어판 = 새 안드로이드 앱이 아니라 '텔레그램 2-way'(업계표준 Freqtrade 패턴). Rauto엔 이미 alert_telegram(발신)"
  "+telegram_poll(수신) 기반 존재. 비상 = 2단 대응(자동 1차 + 사람 2차).", True)
BULL("Binance 선물 서브계정 — 출금 OFF · 거래전용 키 · IP 화이트리스트(본계좌 격리).")
BULL("8 Bots(슬롯) — 전략 플러그인이 신호 발생(§8 봇 무수정, plugin/래퍼).")
BULL("주문 모듈(신규, 체크리스트 B) — 신호→주문: 진입=지정가(메이커, 7h/8h라 100%체결 검증) / 청산=시장가 · "
     "수량 stepSize·최소명목 반올림 · 멱등 클라이언트오더ID · 429/거부 백오프.")
BULL("★Risk Guard(자동 1차, <1초, 엔진 밖 독립레이어) — 사람 개입 전 기계가 먼저 막음.")
BULL("모니터링·알람 — 8슬롯 상태(포지션·미실현손익·슬리피지·MDD) 실시간 집계 + 진입/청산/비상 즉시 발신.")
BULL("Android·Telegram(2차 사람) — /status(8슬롯)·/killall(전부 시장가 청산)·/flat<슬롯>·/pause·/resume.")

H("3. ★2단 비상대응 (캡틴 요구 핵심)")
L("1차(기계, <1초): 슬리피지 브레이크(진입 체결가가 신호가보다 X bp 이상 나쁘면 즉시 시장가 강제청산) · "
  "급변 브레이크(진입 직후 Z초 내 Y% 역행 시 청산) · MDD 하드스톱(계좌 피크 −20% → 전슬롯 청산+정지) · "
  "데이터끊김 가드(피드 N분 정지 시 신규진입 금지).", True)
L("2차(사람, 수초): 텔레그램 알람 수신 → 폰에서 /killall·/flat 즉시 개입. 새벽엔 1차가 먼저 막고 사람이 마무리.", True)
L("[임계값은 테스트넷 실측으로 확정 — X bp·Y%·Z초·N분은 추정 금지, 측정 후 기입]")

H("4. 8슬롯 구성 (캡틴 확정: ②전략추가 + ③기존+성급 동시)")
L("8슬롯 = 다전략 × 다변종 동시 포트폴리오. 합산(8슬롯) 기준으로 노출상한·MDD·모니터링을 잡는다.", True)
L("제안 슬롯맵(확인필요): ③ 기존+성급 동시 = TS인내 / TS성급 / SW참을성(공유) — 성급은 추세봇만(SW성급은 역효과 확인). "
  "② 전략추가 = 확정될 신규 알파 2~3종(예: 성급 확장연구 산물)을 슬롯에 편입. 미확정 슬롯은 '예약'으로 비워둠.")
L("★합산 노출 가드: 8슬롯 각자 레버 → 총명목 폭주 위험. 전역 노출상한(예: 합산 5.6x 캡, k배분) + 계좌단위 MDD −20% 동시충족.", True)

H("5. 완료 조건 (단계 게이트)")
BULL("설계 확정: 본 Work Order + 8슬롯 맵·임계값 변수표 캡틴 승인.")
BULL("테스트넷 실사격: 주문모듈+Risk Guard+Telegram 2-way를 Binance 선물 테스트넷에 올려 "
     "슬리피지 브레이크·/killall·/status 작동 + 실슬리피지 로깅 검증.")
BULL("전환 게이트(체크리스트 19): 동치 7/7 · ★긴급 0 · 슬리피지 p50 ≤ 가정 · 주간 P&L 부호 일치 → 소액 실거래($500~1k).")
BULL("롤백 기준(체크리스트 20): 실계좌 2주 내 MDD −10% 또는 모델괴리 누적 50bp → 페이퍼 복귀.")

H("6. 재평가 시점")
L("착수 = 06-19 페이퍼 1주 §5 보고 + 성급TS 채택 결정 직후. 라이브 3개월 후 듀얼 k·8슬롯 배분 재평가(기존 Work Order DualKUpReeval 연계).")

H("7. 정직 플래그 — 4일 현실 + 미결")
BULL("[정직] 4일 안에 8슬롯 실돈 자동봇+자동브레이크+폰제어를 '안전하게' 런칭 불가. 4일 = 설계확정 + 테스트넷 착수. 실돈은 가드 검증 후.")
BULL("[확인필요] 투입 자본 규모 · 슬롯별 자본배분(공유마진 vs 슬롯별) · ②전략추가의 실제 알파 정의(미발굴분 예약).")
BULL("[확인필요] 임계값(X bp 슬리피지/Y% 급변/Z초/N분) — 테스트넷 실측으로 확정.")
BULL("[신뢰도] 아키텍처 방향=95(업계표준·기존부품 존재) / 임계값·8슬롯 세부=15(미측정, 테스트넷 후).")

H("8. 참조")
L("이번 챗 아키텍처 그래픽(rauto_production_architecture) · Checklist_LiveTransition_Draft_20260612 · "
  "Freqtrade telegram-usage(/status·/fx·order_types emergency_exit→market) · CLAUDE.md §0(MDD−20%)·§7(비용2레이어)·§13.")

d.save(OUT)
print("Work Order saved:", OUT, "| paras", len(d.paragraphs))

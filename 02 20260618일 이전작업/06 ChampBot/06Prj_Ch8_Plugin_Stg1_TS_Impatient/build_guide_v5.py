# -*- coding: utf-8 -*-
# Guide_AlphaDiscovery_Method_v4 → v5: 기존 전체 보존 + 06Prj_Ch8 TIL 7건 추가.
import docx, os, shutil
from docx.shared import Pt, RGBColor

SRC = r"D:\ML\Verify\00WorkHstr\00Basic_Setup_Package\Guide_AlphaDiscovery_Method_v4.docx"
OUT = r"D:\ML\Verify\06Prj_Ch8_Plugin_Stg1_TS_Impatient\docs\Guide_AlphaDiscovery_Method_v5.docx"

d = docx.Document(SRC)


def h(t, sz=15, c=(31, 56, 104)):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(sz)
    r.font.color.rgb = RGBColor(*c); return p


def line(t, bold=False):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = bold; r.font.size = Pt(10.5); return p


d.add_page_break()
h("v5 추가 — 06Prj_Ch8 Plugin (TrendStack Impatient) TIL 7건", 16)
line("출처: 06Prj_Ch8_Plugin_Stg1_TS_Impatient (2026-06-15). 신뢰도 태그: 95=측정/55=추론/15=가설.", False)
line("아래는 '성급(Impatient)' 분기 연구에서 건진 방법론 밧줄들. Part 2 수렁↔밧줄에 준해 정리.", False)

h("TIL-1 ★봇 성격별 '참을성' 정반대 법칙 (신뢰95) — 가장 중요", 13, (192, 0, 0))
line("수렁: 한 봇에서 통한 진입타이밍(성급/인내)을 다른 봇에 그대로 복사.", False)
line("밧줄: 봇의 '성격'으로 진입타이밍을 정하라. 추세추종=성급(피벗대기 제거)이 약 / 평균회귀=인내(바닥확인)가 약.", True)
line("사례: 성급 TS는 전 연도·장세·롱숏 우위(추세 초반 포착). 같은 '성급'을 SidewayDCA(평균회귀)에 적용하니 2026 PF 0.77 손실='떨어지는 칼 잡기'. 피벗확정은 추세봇엔 '늦은 진입', 회귀봇엔 '안전장치'.", False)
line("★향후 확장 가능성(신뢰15, 추가 알파 연구 필요): '성급'은 진입 1축만 바꿔 큰 개선. 같은 원리를 ①부분익절/피라미딩 ②동방향 재진입 ③멀티TF 동시플립 ④강신호(ER/ADX) 게이팅으로 확장하면 추가 알파 여지. 단 추세추종 계열에만, 측정·CPCV 통과 후 채택.", True)

h("TIL-2 '격차가 진짜인가' = 비용 민감도 스윕 (신뢰95)", 13)
line("수렁: 단일 비용 한 점에서 'A가 B보다 좋다' 단정 → 비용/거래빈도 아티팩트일 수 있음.", False)
line("밧줄: 비용을 4→30bp 스윕해 전 구간 우위 유지되면 진짜. + '거래당 평균 net R'(빈도 무관)로 '물량 우위 vs 질 우위' 구분.", True)
line("사례: 성급TS는 30bp에서도 우위. 단 연도별 PF는 평상시 기존이 높고 2025(기존 붕괴) 덕에 전체 우위 — '물량+위기내성'이 정체였음.", False)

h("TIL-3 장봉(7h/8h) 지정가 체결 가능성 검증법 (신뢰95)", 13)
line("수렁: 백테 'close 체결'이 비현실적이라 가정하고 과한 슬리피지를 얹음.", False)
line("밧줄: 다음 1봉(7~8h) 동안 1분봉 고저가 신호가를 터치하는지로 지정가 체결률 측정. 긴 봉이라 100% 체결 → close 체결은 지정가로 현실적, 비용은 메이커(저렴).", True)
line("사례: TS 716건·SW 93건 모두 100% 체결. ★욕심(passive offset 5~20bp)은 역선택 — 놓치는 게 하필 수익 거래(평균 +2.6%).", False)

h("TIL-4 과최적합 2중 점검 = OOS + CPCV 표준6 (신뢰90)", 13)
line("수렁: 전표본 최적화가 -20% 경계에 딱 붙어 짜맞춤(경계 hugging).", False)
line("밧줄: ①학습기간 최적→검증기간 적용(OOS) ②CPCV 표준6(15경로) p25·최악경로 둘 다 >0. 전표본 최적은 '예시'일 뿐 채택불가.", True)
line("사례: 4bp 전표본 최적이 -19.5%(경계). 현실 14bp로 가니 최적 k가 1.05→0.85로 내려가 -17.9%(버퍼). OOS·CPCV 둘 다 통과.", False)

h("TIL-5 현실 비용 = 체결방식별 분해 (신뢰95)", 13)
line("수렁: 왕복 비용을 일률 한 숫자로 가정.", False)
line("밧줄: 체결방식별로 분해. 진입 지정가=메이커(~2bp) / SL청산=스톱→시장가=테이커(~6bp). SL은 메이커 불가. §7 2레이어(신호 4bp는 P&L 금지)와 별개.", True)
line("사례: 양측 테이커 14bp면 TS단독 MDD -20.8%(위반)이나, 진입지정가+청산시장가 ~8bp면 -18.3%(이내).", False)

h("TIL-6 빠른 충실 백테 = 검증된 _step 재사용 + 사후 사이징 (신뢰95)", 13)
line("수렁: 라이브 on_bar 전체 리플레이(1.58M봉)는 느려서 반복 실험 어려움.", False)
line("밧줄: 동치(live≡replay) 검증된 _step의 replay로 거래만 빠르게 생성 → 사이징/P&L은 사후 적용(검증된 resolve_replay). 단 SW처럼 replay 없는 봇은 on_bar 1회 후 결과 csv 재사용.", True)

h("TIL-7 아티팩트 경계 = 평노출×raw 신호R = 거짓 거대수익 (신뢰95)", 13)
line("수렁: raw 신호R에 평균 노출 곱해 복리 → +14461% 같은 거짓 대박.", False)
line("밧줄: 반드시 '확정 베이스라인 재현 대조'로 검증. 절대값이 §9 등 확정치와 자릿수 다르면 아티팩트 의심. MDD가 -20% 크게 위반하면 노출기준 오류 신호.", True)

# 변경 이력 한 줄
line("")
h("[변경 이력] v4 → v5", 12)
line("2026-06-15 · 06Prj_Ch8 · TIL 7건 추가(성급 분기 연구). 핵심=봇 성격별 참을성 정반대 법칙. 출처=06Prj_Ch8_Plugin_Stg1_TS_Impatient.", False)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
d.save(OUT)
print("Guide v5 saved:", OUT, "| paras now", len(d.paragraphs))

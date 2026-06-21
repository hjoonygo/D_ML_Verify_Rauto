# -*- coding: utf-8 -*-
# [build_ControlApp_package.py] 07Prj_Rauto_Phone_Stg4_ControlApp — 폰 제어앱(PWA) 인수인계 redo(새 네이밍·b25·현경로).
import os, glob, zipfile
from docx import Document
V = r"D:\ML\Verify"
M06 = os.path.join(V, "02 20260618일 이전작업", "06 ChampBot", "06Prj_Ch8_Plugin_Stg1_TS_Impatient")
UI = os.path.join(M06, "control_ui")
BASE = os.path.join(V, "00WorkHstr", "00Basic_Setup_Package")
OUTDIR = os.path.join(V, "07 Rauto"); os.makedirs(OUTDIR, exist_ok=True)
NAME = "07Prj_Rauto_Phone_Stg4_ControlApp"
OUT = os.path.join(OUTDIR, NAME + ".zip")
HANDOVER = os.path.join(BASE, "Handover_Rauto_Phone_ControlApp_20260619.docx")


def mkdoc(path):
    d = Document()
    def H(t, l=0): d.add_heading(t, level=l)
    def P(t): d.add_paragraph(t)
    def T(rows, hdr):
        tb = d.add_table(rows=1, cols=len(hdr)); tb.style = 'Light Grid Accent 1'
        for i, h in enumerate(hdr): tb.rows[0].cells[i].text = h
        for r in rows:
            c = tb.add_row().cells
            for i, v in enumerate(r): c[i].text = str(v)
    H("Handover — Rauto 폰 제어앱(PWA) ControlApp", 0)
    P("CCproject=Rauto_Phone · Stg4 · 2026-06-19. 구 Handover_06Prj_Ch8_RautoControlPWA_Chart(2026-06-17)는 폴더이동·b25·새네이밍 미반영=대체됨(보존). 본 문서가 최신.")
    H("1. Output of Chat", 1)
    P("• 안드로이드 폰에 설치하는 Rauto 제어앱(PWA): 자동매매봇 4슬롯(R1~R4)을 폰에서 실시간 보고 통제. RDP 불요.")
    P("• 차트 마커 시각보정 종결(b25): 진입십자·체결마커가 모든 TF(15m/1H/4H)서 캔들에 정확히 붙음.")
    H("2. 구성 (3축)", 1)
    T([["PWA 셸", "manifest.json + sw.js(v25) + icon", "폰 홈에 설치·오프라인 셸·캐시버전으로 강제갱신"],
       ["서버", "control_server.py", "stdlib HTTP. state.json 집계(C:\\Rauto*). git auto-pull(180s)→C:\\Rauto* 자동배포"],
       ["대시보드", "control_dashboard.html(b25)", "자체 캔들차트(TF/EMA/십자선/ㄱ자마커)·4슬롯·챔피언·긴급버튼·실시간가·체결품질"]],
      ["축", "파일", "역할"])
    H("3. 접속·배포", 1)
    P("· 접속: Tailscale serve(사설 HTTPS, https://<host>.ts.net). funnel(공개) 금지·serve만(보안 §control_server).")
    P("· 배포: git push origin master → AWS control_server가 180s마다 pull→C:\\Rauto*(러너)·대시보드 자동반영. 폰은 앱 재실행 시 sw 캐시버전(v25)으로 갱신.")
    H("4. 차트 b25 (이번 핵심 수정)", 1)
    P("진입십자(open_et)·듀얼 닫힌마커(trd)가 7H봉 '시작' 라벨이라 15m/1H서 캔들과 어긋남(4H는 넓어 가려짐). "
      "→ 닫힌거래서 검증된 _fillms(실제 체결 분)으로 통일 + 듀얼 entry 표시 + 15m x축 hh:mm. 검증: 3봇 예외0·진입가가 캔들[저~고] 포함. 커밋 3fc0e40.")
    H("5. 슬롯 표시", 1)
    P("4슬롯 카드(R1 성급·R2 성급왕★챔피언·R3 최적듀얼·R4 최고Calmar듀얼). 챔피언 트로피 2줄(진입가/현재가·실거래OFF·수익률). "
      "실거래 미연결=톤다운. 장세별 PF·연도·롱숏 비교존. (봇 전략·확정수치·검증=07Prj_Rauto_Phone_Stg3_R1234Bot 패키지 참조.)")
    H("6. 파일·현재 경로 (2026-06-18 reorg 후)", 1)
    P("제어앱 원본: D:\\ML\\Verify\\02 20260618일 이전작업\\06 ChampBot\\06Prj_Ch8_Plugin_Stg1_TS_Impatient\\control_ui\\ (git에도 정본).")
    P("배포처: C:\\Rauto1~4(러너·state.json) + control_server가 서빙. 본 zip의 control_app/·live_runners/·engines/에 현 b25 일습 동봉.")
    H("7. 실행(재현)", 1)
    P("① 서버: start_control_gitpull.bat(RAUTO_GIT_PULL=1·RAUTO_REPO=클론경로) 실행 → 4슬롯 러너 자동동기·실행+대시보드 서빙.")
    P("② 폰: Tailscale 켜고 https://<host>.ts.net 접속→홈에 추가(PWA 설치). ③ 갱신 안되면 앱 닫았다 재실행(sw v25).")
    H("8. 네이밍 규약(신규)", 1)
    P("(Proj)Prj_(CCproject)_Stg(결과물횟수)_(작업명). 이 패키지=07Prj_Rauto_Phone_Stg4_ControlApp. 저장 D:\\ML\\Verify\\07 Rauto + G백업. (CLAUDE.md §16 / Guide_FileFolderNaming)")
    d.save(path)


mkdoc(HANDOVER)
CONTROL = [c for c in sorted(glob.glob(os.path.join(UI, "*"))) if os.path.isfile(c) and not c.endswith("state.json")]
RUNNERS = [os.path.join(M06, "rauto1", "test_Rauto1.py"), os.path.join(M06, "rauto2", "test_Rauto2.py"), os.path.join(M06, "rauto3", "test_dual_runner.py")]
ENGINES = sorted(glob.glob(os.path.join(M06, "rauto3", "bots", "*.py")))
DOCS = [HANDOVER, os.path.join(V, "CLAUDE.md"), os.path.join(BASE, "Guide_FileFolderNaming_20260619.docx")]
README = ("07Prj_Rauto_Phone_Stg4_ControlApp — Rauto 폰 제어앱(PWA) 인수인계 (2026-06-19)\n"
          "================================================================\n"
          "안드로이드 폰에서 자동매매봇 4슬롯을 보고 통제하는 PWA 제어앱. b25 차트마커 시각보정 반영.\n"
          "구 Handover_06Prj_Ch8_RautoControlPWA_Chart(06-17)는 폴더이동·b25·네이밍 미반영=대체됨.\n\n"
          "[폴더] control_app/(PWA·서버·대시보드 b25/v25·셋업) · live_runners/(R1~4 러너 b25) · engines/(15) · docs/(인수인계·네이밍·CLAUDE.md).\n"
          "[실행] control_app/start_control_gitpull.bat → 서버+4슬롯 자동 / 폰서 Tailscale HTTPS 접속→PWA 설치.\n"
          "[배포] git push→AWS auto-pull(180s)→C:\\Rauto* 자동반영. 폰 갱신=앱 재실행(sw v25).\n"
          "[봇 전략·확정수치·검증] = 07Prj_Rauto_Phone_Stg3_R1234Bot 패키지 참조.\n")


def add(z, files, arc):
    n = 0
    for f in files:
        if f and os.path.isfile(f) and "__pycache__" not in f and not f.endswith(".pyc"):
            z.write(f, arc + "/" + os.path.basename(f)); n += 1
    return n


cnt = {}
with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as z:
    z.writestr("README_MANIFEST.txt", README)
    cnt["control_app"] = add(z, CONTROL, "control_app")
    cnt["live_runners"] = add(z, RUNNERS, "live_runners")
    cnt["engines"] = add(z, ENGINES, "engines")
    cnt["docs"] = add(z, DOCS, "docs")
print("ZIP:", OUT, "|", round(os.path.getsize(OUT) / 1024, 1), "KB")
for k, v in cnt.items(): print(f"  {k}: {v}")
print("총:", sum(cnt.values()))

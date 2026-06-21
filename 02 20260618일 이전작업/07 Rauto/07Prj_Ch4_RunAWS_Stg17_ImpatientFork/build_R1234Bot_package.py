# -*- coding: utf-8 -*-
# [build_R1234Bot_package.py] 07Prj_Rauto_Phone_Stg3_R1234Bot 패키지 빌드 — 인수인계docx+네이밍docx+전체zip.
import os, glob, zipfile
from docx import Document
V = r"D:\ML\Verify"
M06 = os.path.join(V, "02 20260618일 이전작업", "06 ChampBot", "06Prj_Ch8_Plugin_Stg1_TS_Impatient")
M07 = os.path.join(V, "02 20260618일 이전작업", "07 Rauto", "07Prj_Ch4_RunAWS_Stg17_ImpatientFork")
ML = r"D:\ML"
BASE = os.path.join(V, "00WorkHstr", "00Basic_Setup_Package")
OUTDIR = os.path.join(V, "07 Rauto"); os.makedirs(OUTDIR, exist_ok=True)
NAME = "07Prj_Rauto_Phone_Stg3_R1234Bot"
OUT = os.path.join(OUTDIR, NAME + ".zip")


def mkdoc_handover(path):
    d = Document()

    def H(t, l=0): d.add_heading(t, level=l)
    def P(t): d.add_paragraph(t)

    def T(rows, hdr):
        tb = d.add_table(rows=1, cols=len(hdr)); tb.style = 'Light Grid Accent 1'
        for i, h in enumerate(hdr): tb.rows[0].cells[i].text = h
        for r in rows:
            c = tb.add_row().cells
            for i, v in enumerate(r): c[i].text = str(v)
    H("Handover — 4봇 실슬리피지 백테 ↔ Rauto 시스템 연동 (R1234Bot)", 0)
    P("CCproject=Rauto_Phone · Stg3 · 2026-06-19 · 봇별 Plugin 단일출처 + 재현용(검증데이터+검증PY) 일괄.")
    H("1. Output of Chat", 1)
    P("• 4봇(R1성급·R2성급왕·R3최적듀얼·R4최고Calmar듀얼)을 검증엔진 무수정 on_bar(핀고정)+실슬립(5bp)로 36개월 확정.")
    P("• 각 봇을 Plugin_06ChmpBot_(봇명).py 단일출처로 통합(설정+확정수치+재현함수+Rauto 슬롯 연동). 단독 실행으로 숫자 재현.")
    P("• Rauto 라이브 차트 b25 수정(open_et·듀얼 trd=체결분 _fillms, 전 TF 캔들정렬·듀얼 entry 표시) 반영.")
    H("2. 4봇 ↔ Rauto 슬롯 매핑 + 확정수치", 1)
    T([["R1 성급(ImpatientTS)", "test_Rauto1.py / C:\\Rauto1", "단독$10k·lev22", "+5,932%/-20.0%/PF1.72/승34", "Plugin_..._R1_ImpatientTS"],
       ["R2 성급왕(ImpatientKingTS)★챔피언", "test_Rauto2.py / C:\\Rauto2", "단독$10k·lev22", "+11,397%/-17.3%/PF1.90/승34", "Plugin_..._R2_ImpatientKingTS"],
       ["R3 최적듀얼", "test_dual_runner(R3)/C:\\Rauto3", "듀얼$20k·k1.1·ER0.4", "+8,850%/-18.0%/PF1.94/승38", "Plugin_..._R3_OptimalDual"],
       ["R4 최고Calmar듀얼", "test_dual_runner(R4)/C:\\Rauto4", "듀얼$20k·k1.4·ER0.4", "+30,156%/-23.4%(위반)/PF1.94", "Plugin_..._R4_MaxCalmarDual"]],
      ["봇", "라이브 슬롯", "설정", "확정 36mo(실슬립)", "Plugin"])
    P("전봇 매년양수(2023~2026)·롱숏 양쪽수익(숏 약간우세). R4는 MDD-20% 위반=실거래 부적합(R2/R3 권장).")
    H("3. 실슬립 백테 방법 (백테 5관문 = CLAUDE.md §15)", 1)
    P("①검증엔진만(재구현 금지) ②앵커/동치 대조(성급batch716=measure_slippage·성급핀+5932≈앵커+5791) "
      "③손절터치 윈도우=et+7H ④슬립명시(진입/청산~0bp, 손절5bp=검증0~20bp 견고) ⑤기준명시($10k단독/$20k듀얼).")
    P("핀고정 on_bar: 성급에 GRID_ANCHOR 래퍼로 7H그리드 정렬→백테 resample과 일치. 산출=led36_{king,imp_pinned}.csv.")
    H("4. b25 차트 수정 (라이브 연동)", 1)
    P("진입십자(open_et)·듀얼 닫힌마커(trd)가 7H봉 시작라벨이라 15m/1H서 캔들과 어긋남(4H는 가려짐). "
      "→ 닫힌거래서 검증된 _fillms(체결분)으로 통일+듀얼 entry 표시+15m x축 hh:mm. 검증: 3봇 예외0·진입가 캔들[저~고] 포함. push 3fc0e40.")
    H("5. 재현 절차 (★검증데이터+검증PY 동봉 필수)", 1)
    P("빠른재현: python Plugin_06ChmpBot_(봇).py → 확정수치 출력(입력=led36_*.csv·sw_patient_er.csv, Merged_Data 불요).")
    P("깊은재현(원장 재생성): bt36_ledgers.py + Merged_Data.csv(455MB, 06_Data DATA_NOTE의 위치·생성법). "
      "Merged_Data 생성 = 02_DataProcessing(merge_monthly+compute_oi_derived+merge_oi_metrics).")
    H("6. zip 구성", 1)
    P("00_Plugins(4봇+plugin_common) / 01_Engines(15 무수정) / 02_LiveRunners(러너 b25+제어앱) / "
      "03_DataProcessing(원천+어댑터) / 04_Tests / 05_Verification(검증가공+engines) / 06_Data(원장+ER+노트) / 07_Docs.")
    H("7. 다음 단계", 1)
    P("① 실거래 소액(R2 챔피언·출금권한없는키·상주주문봇=Work_Order_RautoLiveTransition) "
      "② 틱정밀화(봇계측=Work_Order_TickSlippage_Precision) ③ R1 그리드 핀.")
    d.save(path)


def mkdoc_naming(path):
    d = Document()

    def H(t, l=0): d.add_heading(t, level=l)
    def P(t): d.add_paragraph(t)
    H("파일·폴더 네이밍 관리지침 (2026-06-19)", 0)
    P("출처: 캡틴 지시 2026-06-19. 단일출처=CLAUDE.md §16. 본 문서는 그 사본·설명.")
    H("1. 핵심 변경 — CH 폐기", 1)
    P("CH(채팅회차)명은 의미 없어져 폐기. 의미 단위 = (Proj번호)Prj + CCproject명 + Stg(결과물 작업횟수)_작업명.")
    H("2. 구성요소", 1)
    P("· CCproject = 이 ClaudeCode 작업 프로젝트명(캡틴 지정). 예: Rauto_Phone.")
    P("· Stg횟수 = '결과물이 나온 작업' 횟수(AI가 세서 정함).")
    P("· 작업명 = 짧고 의미있게(미지정 시 AI가 정함). 예: R1234Bot.")
    H("3. 명명·저장 규칙", 1)
    P("산출물명 = (Proj번호)Prj_(CCproject)_Stg(횟수)_(작업명).")
    P("예: 07Prj_Rauto_Phone_Stg3_R1234Bot")
    P("저장 = D:\\ML\\Verify\\07 Rauto\\(풀네임).zip · 백업 = G:\\내 드라이브\\00AI개발지식DB\\자산관리\\유동자산\\자동매매\\07 Rauto\\(풀네임).zip")
    P("다른 Proj(04~08)도 동일 패턴(해당 NN 폴더).")
    H("4. ★재현 필수 원칙", 1)
    P("zip엔 '검증 데이터 + 검증한 PY'를 반드시 함께 담는다 — 둘 다 있어야 재현 가능. 봇별 단일출처=Plugin_06ChmpBot_(봇명).py.")
    H("5. ★2026-06-18 이전작업 보관 위치", 1)
    P("두 폴더에 몰아넣음(옛 파일은 여기서 찾는다):")
    P("· D:\\ML\\Verify\\02 20260618일 이전작업\\")
    P("· G:\\내 드라이브\\00AI개발지식DB\\자산관리\\유동자산\\자동매매\\02 20260618일 이전작업\\")
    P("내부 구조: 03 InfraA_Mod · 04 IDEA4Concept · 05 Alpha_Up · 06 ChampBot · 07 Rauto (기존 유지). 옛 06Prj_Ch*·07Prj_Ch* 전부 여기.")
    H("6. 나머지", 1)
    P("폴더 3종(test_/check_/run.bat)·check.py 3역할·INDEX 한 줄·문서맵 등은 기존 CLAUDE.md §4·§13과 동일.")
    d.save(path)


HANDOVER = os.path.join(BASE, "Handover_Rauto_Phone_R1234Bot_20260619.docx")
NAMINGDOC = os.path.join(BASE, "Guide_FileFolderNaming_20260619.docx")
mkdoc_handover(HANDOVER)
mkdoc_naming(NAMINGDOC)

PLUGINS = sorted(glob.glob(os.path.join(M07, "Plugin_06ChmpBot_*.py"))) + [os.path.join(M07, "plugin_common.py")]
ENGINES = sorted(glob.glob(os.path.join(M06, "rauto3", "bots", "*.py")))
RUNNERS = [os.path.join(M06, "rauto1", "test_Rauto1.py"), os.path.join(M06, "rauto2", "test_Rauto2.py"), os.path.join(M06, "rauto3", "test_dual_runner.py")]
CONTROL = [c for c in sorted(glob.glob(os.path.join(M06, "control_ui", "*"))) if not c.endswith("state.json") and os.path.isfile(c)]
DP8 = os.path.join(V, "08 BTC_Data")
DATAPROC = [os.path.join(DP8, "08Prj_Dauto_Ch2_Stg13_OiZscoreAdapter", x) for x in ("compute_oi_derived_features.py", "dauto_loader.py", "oi_zscore_adapter.py")] + \
           [os.path.join(DP8, "08Prj_Dauto_Ch1_Collector_Stg1_RestPoller", x) for x in ("dauto_collector.py", "report_daily_coverage.py")] + \
           [os.path.join(M06, "rauto3", "bots", a) for a in ("atr_ratio_adapter.py", "regime_feature_extractor.py")]
ALL7 = sorted(glob.glob(os.path.join(M07, "*.py")))
SKIP = set(os.path.basename(p) for p in PLUGINS)
TESTPRE = ("test_", "check_", "diag_", "validate_", "verify_", "graph_")
TESTS = [f for f in ALL7 if os.path.basename(f).startswith(TESTPRE) and os.path.basename(f) not in SKIP]
VERIF = [f for f in ALL7 if not os.path.basename(f).startswith(TESTPRE) and os.path.basename(f) not in SKIP]
VERIF_ENG = sorted(glob.glob(os.path.join(M07, "bots", "*.py")))
DATA = [os.path.join(M07, x) for x in ("led36_king.csv", "led36_imp_pinned.csv", "sw_patient.csv", "sw_patient_er.csv", "comprehensive_4bot.json")]
DOCS = [os.path.join(V, "CLAUDE.md"), HANDOVER, NAMINGDOC] + [os.path.join(BASE, x) for x in (
    "Work_Order_TickSlippage_Precision_20260618.docx", "Work_Order_RautoLiveTransition_20260620.docx")]

DATANOTE = ("[Merged_Data.csv — 455MB, zip 미포함]\n위치: D:\\ML\\Verify\\Merged_Data.csv (1m 2023-05-01~2026-04-30, 157.8만행, OHLCV+oi_zscore_24h)\n"
            "생성: 02_DataProcessing의 merge_monthly_csv_V8X(월별1m병합)+compute_oi_derived_features(OI파생)+merge_oi_metrics_36mo(병합).\n"
            "oi_zscore_24h=REPAIRED 계보(CLAUDE.md §8). 원천 1m=Binance Vision(과거) / C:\\BinanceData(라이브).\n"
            "빠른재현은 Merged 불요(led36_*.csv·sw_patient_er.csv로 Plugin 실행). 깊은재현(원장 재생성)에만 필요.\n"
            "데이터가공 코드 위치(2026-06-18 reorg 후): 03_DataProcessing = D:\\ML\\Verify\\08 BTC_Data\\08Prj_Dauto_*\n"
            "  (compute_oi_derived_features·dauto_collector·dauto_loader·oi_zscore_adapter + 런타임 atr/regime 어댑터).\n"
            "  ★월별 1m→Merged 병합 스크립트(구 D:\\ML\\merge_monthly_csv_V8X·merge_oi_metrics_36mo)는 reorg로 현재 미발견 — 복구 필요시 백업/구드라이브 확인.")
README = ("07Prj_Rauto_Phone_Stg3_R1234Bot — Rauto 4봇 일괄 패키지 (2026-06-19)\n"
          "================================================================\n"
          "CCproject=Rauto_Phone · Stg3 · 작업명=R1234Bot. 명명·저장 규칙=CLAUDE.md §16 / 07_Docs/Guide_FileFolderNaming.\n\n"
          "[폴더] 00_Plugins(4봇 단일출처+plugin_common) · 01_Engines(15 무수정 §8) · 02_LiveRunners(러너 b25+제어앱) ·\n"
          " 03_DataProcessing(merge/oi 원천+어댑터) · 04_Tests · 05_Verification(검증가공+engines) · 06_Data(원장+ER+노트) · 07_Docs.\n\n"
          "[빠른재현] python 00_Plugins/Plugin_06ChmpBot_(봇).py → 확정수치. (Merged_Data 불요)\n"
          "[확정 4봇] R1+5,932%/-20.0% · R2성급왕+11,397%/-17.3%(챔피언) · R3듀얼k1.1+8,850%/-18.0% · R4듀얼k1.4+30,156%/-23.4%(MDD위반).\n"
          "[★b25] 진입십자·마커 체결분 정렬(open_et/trd=_fillms). [비용] 실P&L 8bp(메이커2+테이커4+슬립2)+펀딩(§7).\n")


def add(z, files, arc):
    n = 0
    for f in files:
        if f and os.path.isfile(f) and "__pycache__" not in f and not f.endswith(".pyc"):
            z.write(f, arc + "/" + os.path.basename(f)); n += 1
    return n


cnt = {}
with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as z:
    z.writestr("README_MANIFEST.txt", README)
    cnt["00_Plugins"] = add(z, PLUGINS, "00_Plugins_4봇")
    cnt["01_Engines"] = add(z, ENGINES, "01_Engines_검증엔진")
    cnt["02_Runners"] = add(z, RUNNERS, "02_LiveRunners/runners") + add(z, CONTROL, "02_LiveRunners/control")
    cnt["03_DataProc"] = add(z, DATAPROC, "03_DataProcessing_데이터가공")
    cnt["04_Tests"] = add(z, TESTS, "04_Tests_테스트")
    cnt["05_Verif"] = add(z, VERIF, "05_Verification_검증가공") + add(z, VERIF_ENG, "05_Verification_검증가공/engines")
    cnt["06_Data"] = add(z, DATA, "06_Data_데이터")
    z.writestr("06_Data_데이터/DATA_NOTE.txt", DATANOTE)
    cnt["07_Docs"] = add(z, DOCS, "07_Docs_문서")
print("ZIP:", OUT, "|", round(os.path.getsize(OUT) / 1024, 1), "KB")
for k, v in cnt.items(): print(f"  {k}: {v}")
print("총:", sum(cnt.values()))

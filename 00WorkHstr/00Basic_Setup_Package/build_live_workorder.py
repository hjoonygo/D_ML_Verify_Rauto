# -*- coding: utf-8 -*-
# [build_live_workorder.py] Rauto 실거래 전환 WorkOrder + 인수인계보고서 docx 생성 (1회용)
import os, sys
import docx
from docx.shared import Pt
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
PKG = r"D:\ML\Verify\00WorkHstr\00Basic_Setup_Package"


def add(doc, style, text):
    if style == "h1": doc.add_heading(text, level=1)
    elif style == "h2": doc.add_heading(text, level=2)
    elif style == "code":
        p = doc.add_paragraph(); r = p.add_run(text); r.font.name = "Consolas"; r.font.size = Pt(8)
    else: doc.add_paragraph(text)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers): t.rows[0].cells[i].text = h
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row): c[i].text = str(v)


# ============================ WORK ORDER ============================
def build_workorder():
    doc = docx.Document()
    add(doc, "h1", "Work Order — Rauto 실거래(소액) 전환 검증·보강")
    add(doc, "p", "작성 2026-06-20 | 작성자 Claude(Opus 4.8, Claude Code) | src=캡틴 지시 2026-06-20(6/20 페이퍼 테스트 후 소액 실거래 단계)")
    add(doc, "p", "한 줄: 06-20 페이퍼 테스트 통과 → 소액 실거래로 간다. 단 ★현 시스템엔 '실주문을 내는 코드가 전혀 없다'(신호생성+페이퍼 P&L 시뮬까지). "
                  "실거래는 상주 주문봇을 새로 만들고, 아래 P0 안전장치를 테스트넷·소액으로 통과한 뒤에만 개시한다. 본 문서 = 코드·데이터흐름·타점 감사 + 리스크 체크리스트 + 시나리오별 단계검증·보강 계획.")

    add(doc, "h1", "1. 맥락 / 전제")
    add(doc, "p", "확정 알파(채택): R2 성급왕TS 단독(인트라바 손절+쿨다운, CPCV 표준6 통과 p25+24%). 챔피언=R2. R1 성급단독·R3 최적듀얼·R4 최고Calmar듀얼은 비교/관찰. "
                  "실비용 단일출처(§7 [버전B] 8bp = 진입 메이커2+청산 테이커4+슬립2, 펀딩 별도). 목표(§0): 월+10%·매월 양수·절대 MDD -20% 이내.")
    add(doc, "p", "실거래 자본: 소액부터(예 $100~$1,000), 마이크로 사이즈로 시작. 절대 출금권한 없는 API키.")

    add(doc, "h1", "2. 현 시스템 진단 (코드·데이터흐름·타점 — 차근차근)")
    add(doc, "h2", "2.1 코드 구조")
    add(doc, "p", "구성: Dauto 수집(REST 폴러) → C:\\BinanceData CSV → 봇(신호 on_bar) → rauto_paper_engine(가상 P&L) → state.json → 제어앱(대시보드). "
                  "★rauto_paper_engine.resolve_replay(R, mae, fund)는 '사전계산된 수익 R을 적용'하는 시뮬레이터다(슬리피지·MMR티어·하드스탑 청산 모델 내장하나 전부 가상). "
                  "전수 검색 결과 거래소 주문 코드(ccxt/fapi/create_order/clientOrderId/reduceOnly) 0건 = 실주문 레이어 전무.")
    add(doc, "p", "rauto_contract.py: Signal(action/side/size_pct/leverage/sl/tp), Fill(체결 통지), BotPlugin.on_fill 인터페이스는 존재 → 주문봇이 이 계약을 구현하면 봇 본체(§8 해시락) 무수정으로 결합 가능. "
                  "단 Signal엔 주문타입(지정/시장/스톱) 구분·clientOrderId·멱등키 없음 → 주문봇이 보강.")
    add(doc, "h2", "2.2 데이터 흐름 (현재 vs 실거래 필요)")
    add(doc, "p", "현재: 매시간 '전체 재생(replay)' = 누적 CSV를 처음부터 재관통해 state.json 재작성(멱등·배치). "
                  "실거래 필요: 실시간 피드(웹소켓 권장) → 봉 합성 → 봉마감 신호 → ★실주문 발행 → 체결확인(User Data Stream) → 포지션 대사 → 영속 장부. 이벤트 기반·상주.")
    add(doc, "h2", "2.3 타점 결정 (검증된 부분 + 갭)")
    add(doc, "p", "타점: TS=7H봉·SW=8H봉 '마감 시' 진입/청산 판정, 체결=다음봉 시가(지정가 가정)·손절=시장가. 신호 자체는 인과적(룩어헤드 없음, 이번 세션 3각 감사 확인). "
                  "갭: ①지정가 '다음 긴봉 창 100% 체결' 전제(§7 [B] 근거)가 변동장서 깨짐=최대 갭 ②스톱 시장가 슬리피지/갭 ③'동치' 테스트는 결정성만 검증(인과성 아님) ④Rauto1 봉경계 epoch(백테와 5h 어긋남; R2는 핀고정 완료 399c1b8).")

    add(doc, "h1", "3. ★최대 블로커 — 상주 주문봇(live executor) 부재")
    add(doc, "p", "실거래의 단일 최대 과제. 새로 만들어야 할 컴포넌트: (1)봇 Signal→거래소 주문 변환(지정가 진입·STOP_MARKET reduceOnly 손절) (2)체결확인(웹소켓 User Data Stream 1차+REST 백업) "
                  "(3)멱등 주문(결정적 clientOrderId, 타임아웃 시 Query 후 재전송) (4)포지션 대사(거래소=진실원장) (5)킬스위치 (6)웹소켓 견고성. 페이퍼엔진은 P&L 검증용으로 병행.")

    add(doc, "h1", "4. 실거래 개시 전 필수 체크리스트 (Binance 공식문서 기반, 우선순위)")
    add(doc, "p", "P0=없으면 개시금지 · P1=개시 전 강력권장 · P2=초기운영 중 정착. (출처: developers.binance.com 공식 + 업계합의, 신뢰도 95/55)")
    table(doc, ["#", "우선", "항목", "통과 기준"], [
        ["1", "P0", "킬스위치(일일최대손실+누적MDD)", "도달 시 신규차단+reduceOnly 청산 자동(테스트넷 검증)"],
        ["2", "P0", "API키 권한 최소화", "출금 비활성+IP 화이트리스트+시크릿 격리"],
        ["3", "P0", "포지션모드 명시(hedge/one-way)", "시작시 set&조회. 듀얼 동일심볼 롱숏 동시면 hedge"],
        ["4", "P0", "모든 청산 reduceOnly/closePosition", "의도치 않은 신규·증액 0건"],
        ["5", "P0", "clientOrderId 멱등 재시도", "타임아웃 시 Query후 재전송, 중복주문 0"],
        ["6", "P0", "부팅 대사 + 고아주문 정리", "거래소로 내부장부 재구성·stale 취소"],
        ["7", "P0", "주문 사전검증(tick/step/minQty/minNotional)", "exchangeInfo 라운딩, -1013 거부 0"],
        ["8", "P0", "레버·단일포지션 명목 하드캡", "22 초과·캡 초과 주문 차단"],
        ["9", "P1", "시계동기(NTP+서버오프셋), recvWindow 5000", "-1021 거부 0"],
        ["10", "P1", "웹소켓 재연결(지수백오프)+listenKey 30분 keepalive+재연결후 대사", "강제끊김 테스트 자동복구"],
        ["11", "P1", "지정가 미체결 타임아웃·부분체결 잔량정책", "미체결 취소/전환, 실체결만 장부"],
        ["12", "P1", "슬리피지 가드(시장가 손절 초과시 알림/축소)", "한도초과 감지"],
        ["13", "P1", "청산가 모니터(티어형 MMR)+mark가 거리알람", "진입전 청산버퍼 검증"],
        ["14", "P1", "데이터 워치독(봉갭/지연→신호보류)+봉경계 백테정합", "갭시 진입보류, resample 원점 일치"],
        ["15", "P1", "비용 단일출처(8bp+실펀딩)", "§7 [B] 적용, 이중계산 0"],
        ["16", "P1", "알림 7종(체결·오류·킬·끊김·대사불일치·일손익·펀딩)", "실트리거 수신 확인"],
        ["17", "P1", "페이퍼-라이브 병행+fill 슬립분포 대조", "거래단위 괴리 측정"],
        ["18", "P2", "램프업 게이트+거래소 점검대응", "50~100거래·신고점·괴리충족 시만 증액"],
    ])

    add(doc, "h1", "5. 시나리오별 단계검증 (네 판단 — 다양한 상황 가정)")
    add(doc, "p", "각 시나리오: 가정 → 위험 → 검증방법 → 보강. 테스트넷/소액에서 강제 재현해 통과 확인.")
    table(doc, ["시나리오", "가정", "위험", "검증·보강"], [
        ["S1 미체결", "급변장서 지정가 진입 미체결", "봇은 진입완료로 알지만 실포지션 0 → 유령포지션", "N분 타임아웃→취소/시장전환/스킵. 체결확인 후에만 보유처리. 대사로 즉시 검출"],
        ["S2 부분체결", "지정가 일부만 체결", "수량·평단 불일치, 손절수량 오류", "executedQty만 인정, 잔량 취소. 손절은 closePosition=true"],
        ["S3 스톱 갭", "급락으로 손절가 갭 통과", "손절가 ≪ 체결가, -1% 의도가 -5%", "STOP_MARKET reduceOnly + 슬립가드 알림. 청산버퍼 사전검증(레버22 주의)"],
        ["S4 API 끊김", "웹소켓/REST 단절", "체결·포지션 못 받음, 맹목 거래", "끊김 감지→신규정지(기존보호). 재연결 지수백오프+listenKey keepalive+복귀후 대사"],
        ["S5 봇 재시작", "프로세스 죽었다 부활", "내부상태 소실→이중진입·고아주문", "부팅 대사(거래소=진실)+고아 취소+영속 장부 복구"],
        ["S6 중복주문", "주문 타임아웃 후 재시도", "같은 신호 2회 체결", "결정적 clientOrderId+Query후 재전송(멱등)"],
        ["S7 청산근접", "레버22, mark가 청산가 근접", "강제청산(전액손실급)", "티어형 MMR 청산가 산출+거리알람+포지션 명목 하드캡"],
        ["S8 데이터갭", "Dauto 1분봉/OI/펀딩 누락", "봉 합성오류→오신호", "워치독: 갭/지연시 신호보류+REST 백필+무결성검증"],
        ["S9 봉경계 불일치", "라이브 봉 ≠ 백테 봉", "다른 봉서 신호=백테와 다른 거래", "resample 원점(첫날자정UTC) 핀고정. R1도 R2처럼 적용(현재 미적용)"],
        ["S10 킬스위치", "일손실/MDD 한도 도달", "오발동(좋은 포지션 청산) or 미발동(폭주)", "다층(일손실/MDD/연속에러)+미실현스파이크 구분+쿨다운·수동재개"],
        ["S11 펀딩", "00/08/16 UTC 펀딩 시점 보유", "P&L 과대(펀딩 미반영)", "실펀딩 P&L 반영(§7), 펀딩직전 정책 명문"],
        ["S12 백테-라이브 괴리", "라이브 성능<백테(슬립·미체결)", "과대평가된 알파로 증액", "일일 괴리추적+임계초과시 증액보류. 체결 룩어헤드 교차감사(2f2cd5f) 라이브 적용"],
    ])

    add(doc, "h1", "6. 보강 작업 제안 (단계별 Stg — 테스트넷 우선)")
    add(doc, "p", "원칙: §1 봇 본체·신호엔진(§8 해시) 무수정. 주문봇은 별도 컴포넌트로 신규. 각 Stg는 PASS 후 다음으로.")
    table(doc, ["Stg", "내용", "완료조건"], [
        ["Live-1 주문봇 골격(testnet)", "Signal→지정가진입+STOP_MARKET reduceOnly 손절, exchangeInfo 사전검증, clientOrderId 멱등", "테스트넷서 진입·손절 1회 정상체결, 필터거부 0"],
        ["Live-2 대사·킬·복구", "부팅대사+고아정리+영속장부, 킬스위치(일손실/MDD), 재시작 복구", "S4·S5·S6·S10 강제재현 PASS"],
        ["Live-3 견고성", "웹소켓 User Data Stream+재연결 백오프+keepalive, 데이터 워치독, 시계동기", "강제끊김·갭 재현서 자동복구"],
        ["Live-4 소액 라이브", "실계좌 마이크로 사이즈, 페이퍼 병행, fill 슬립분포·실펀딩·괴리 실측", "P0 전부+P1 핵심 통과, 슬립 분포 측정"],
        ["Live-5 램프업", "50~100거래·신고점·괴리임계 충족 시 단계 증액. 챔피언(R2)만 실거래", "게이트 충족·문서화 후 증액 승인"],
    ])

    add(doc, "h1", "7. 완료조건 / 재평가시점")
    add(doc, "p", "완료조건: ①Live-1~3 테스트넷 PASS(P0 8개 전부) ②Live-4 소액 라이브서 P1 실측(슬립·대사·알림·청산버퍼) ③Live-5 50~100거래 후 라이브-백테 괴리 임계 내 + 신고점 → 단계 증액. "
                  "절대선: 누적 MDD -20% 위반 시 즉시 정지(보수적 -12~15% 1차 경보). 실계좌 = 챔피언(R2)만, 출금권한 없는 키.")
    add(doc, "p", "재평가시점: 각 Live-Stg 종료 후 + 첫 50거래 후 + 월말. CPCV/워크포워드 통과만 채택(§5.7). 라이브 표본이 긴 횡보장을 포함하면 듀얼(R3/R4)의 SW 가치 재평가.")

    add(doc, "h1", "8. 이미 보유한 라이브화 자산 (재활용)")
    add(doc, "p", "봉경계 핀고정(커밋 399c1b8) · 체결가 룩어헤드 교차감사(2f2cd5f) · alert_telegram ops v2 7종(§8) · 비용 단일출처(§7) · CPCV 채택기준(§5.7) · 킬플래그/제어앱 긴급2단 · git auto-pull 무인배포. "
                  "신규로 채울 핵심 = 상주 주문봇의 멱등주문·대사·킬스위치, 웹소켓 견고성, 포지션모드·reduceOnly 정합.")

    out = os.path.join(PKG, "Work_Order_RautoLiveTransition_20260620.docx"); doc.save(out); return out


# ============================ HANDOVER ============================
def build_handover():
    doc = docx.Document()
    add(doc, "h1", "Handover_Rauto_LiveReadiness — 인수인계보고서")
    add(doc, "p", "작성 2026-06-20 | 채팅: Rauto 제어앱·차트·성급왕TS·4슬롯·실거래 전환준비 | 작성자 Claude(Opus 4.8, Claude Code)")
    add(doc, "p", "★한 줄: 페이퍼 단계 완성(4슬롯·R2 챔피언·무인 운영). 실거래 직전. 단 실주문 레이어는 미구축 = 다음 과제(Work_Order_RautoLiveTransition_20260620).")

    add(doc, "h1", "★Output of Chat (다음 채팅 기준점)")
    add(doc, "p", "① 챔피언=R2 성급왕TS 단독(인트라바 손절+쿨다운). CPCV 표준6 통과(p25 +24%·최악 +20% vs 성급), 신규 fitted param 0. 36mo: +8251%/MDD-19.4%/Calmar17.4.")
    add(doc, "p", "② 4슬롯 가동: R1 성급단독 / R2 성급왕(★챔피언) / R3 최적듀얼(k1.1/er0.4/w0) / R4 최고Calmar듀얼(k1.4). 대시보드 합산표시. control_server 자동배포(git auto-pull).")
    add(doc, "p", "③ ★객관 결론: 듀얼(성급왕+인내SW)은 이 36mo서 단독 R2를 못 이김(Calmar 12.7<17.4). 인트라바 손절이 꼬리위험 잡아 SW쿠션 한계효용↓. SW는 '장세보험'(긴 횡보장서 재평가). 역회귀 컷=손해(위험다이얼).")
    add(doc, "p", "④ ★봉경계 핀고정: R2는 백테 resample 그리드(첫날자정UTC 원점)에 앵커 완료(399c1b8, live≡backtest 89~95%). R1은 미적용(epoch, 다음 과제).")
    add(doc, "p", "⑤ ★실거래 전 최대 과제: 상주 주문봇 신규 구축(현재 실주문 코드 0건). Work Order의 P0 체크리스트·시나리오·Live-Stg 참조.")

    add(doc, "h1", "1. 작업 히스토리 (이번 세션)")
    add(doc, "p", "제어앱(PWA): 차트 마커 이탈버그 종결(봇 체결시각 1분 역산 보정 + near()/클램프, b20) · 패닝/Ctrl휠줌/더블탭(b21) · 체결품질 진단패널(b22). git auto-pull로 AWS 무인배포(봇·대시보드).")
    add(doc, "p", "검증: 룩어헤드 3각 감사(미래누출 없음, feature 인과동일·진입청산 인과·P&L 1:1) · ★'동치'=결정성만 검증(인과성 아님) 발견 · 현실체결 A/B(realistic_exec) · 성급왕 CPCV 통과.")
    add(doc, "p", "성급왕TS: 인트라바 손절가드+재진입 쿨다운(naive판 churn 아티팩트 교정) · 봉경계 핀고정 · Rauto2 등록. 4슬롯+R2챔피언. 듀얼 최적조합 스윕(단독이 최적 확인).")
    add(doc, "p", "주요 자체교정(정직): batch 근사오류·슬리피지 22배 단위오류·naive churn — 모두 충실 재시뮬로 적발·수정. 캡틴의 '철저히 크로스체크' 지시가 반복 적중.")

    add(doc, "h1", "2. 시스템 아키텍처")
    add(doc, "p", "데이터: Dauto수집(REST)→C:\\BinanceData CSV→봇 신호(on_bar, 7H/8H봉)→rauto_paper_engine(가상P&L)→state.json→control_server(C:\\Rauto*\\state.json 합산)→PWA 대시보드(Tailscale HTTPS, 폰).")
    add(doc, "p", "슬롯=봇1개(C:\\Rauto1..8). R1 test_Rauto1(성급)·R2 test_Rauto2(성급왕)·R3/R4 test_dual_runner(성급왕+인내SW, env 파라미터). 배포=git push→AWS control_server 180s pull+자동 동기화·실행.")
    add(doc, "p", "봇 본체·신호엔진=§8 해시락 무수정. 변종은 상속 래퍼(성급왕=TrendStackImpatientBot 상속+인트라바가드).")

    add(doc, "h1", "3. 인수인계 주의점 (미검증·점프)")
    add(doc, "p", "[신뢰95] 실주문 레이어 전무 = 실거래 최대 블로커. [신뢰90] R1 봉경계 미핀고정(epoch). [신뢰90] '동치'는 인과성 미검증(determinism only). "
                  "[신뢰55] 듀얼 SW는 이 36mo 저평가 가능(횡보 적었음). [신뢰90] 비용 [B]8bp는 '7H/8H 긴봉 다음창 100%체결' 전제 — 라이브 변동장서 깨질 수 있는 1순위 갭.")

    add(doc, "h1", "4. 다음 채팅 작업계획")
    add(doc, "p", "1순위: Work_Order_RautoLiveTransition_20260620 — 상주 주문봇 Live-1~5(테스트넷→소액→램프업). 2순위: R1 봉경계 핀고정. 3순위: 06-19/20 표본으로 R1~R4 §5 비교. 4순위: SW 라이브 생성 정합·역회귀 듀얼 레버.")

    add(doc, "h1", "5. 신뢰도 점수표")
    add(doc, "p", "95: 성급왕 CPCV 표준6 통과(15경로 p25+24%·전부양수) / 봉경계 핀고정 검증(live≡backtest 89~95%) / 룩어헤드 미래누출 없음(3각 감사) / 실주문 코드 0건(전수검색) / 4슬롯·R2챔피언 합산검증")
    add(doc, "p", "55: 듀얼 SW 36mo 저평가 가능성 / 라이브 체결 슬립 실측치(미측정) / 램프업 임계 구체값")
    add(doc, "p", "15: 향후 긴 횡보장서 듀얼이 단독을 역전할지(미래 의존)")

    out = os.path.join(PKG, "Handover_Rauto_LiveReadiness_20260620.docx"); doc.save(out); return out


if __name__ == "__main__":
    print("workorder:", build_workorder())
    print("handover:", build_handover())

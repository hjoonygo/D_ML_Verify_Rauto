# -*- coding: utf-8 -*-
# [파일명] build_handover_07Prj_Ch4_stg16.py — 인수인계보고서+Key노트 docx 생성기 (1회용)
import os, sys, datetime as dt
import docx
from docx.shared import Pt

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BASE = r"D:\ml\verify"
PKG = os.path.join(BASE, r"00WorkHstr\00Basic_Setup_Package")
S14 = os.path.join(BASE, "07Prj_Ch4_RunAWS_Stg14_LivePaperWarmup")
S16 = os.path.join(BASE, "07Prj_Ch4_RunAWS_Stg16_OpsGuard")


def add(doc, style, text):
    if style == "h1":
        doc.add_heading(text, level=1)
    elif style == "h2":
        doc.add_heading(text, level=2)
    elif style == "code":
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(8)
    else:
        doc.add_paragraph(text)


HANDOVER = [
("h1", "Handover_07Prj_Ch4_RunAWS_stg16 — 인수인계보고서"),
("p", "작성 2026-06-13 | 채팅: 07Prj_Ch4_RunAWS (Stg14 무인자동화 + Stg16 OpsGuard) | 작성자: Claude (Fable 5, Claude Code 세션)"),
("p", "★이번 채팅의 한 줄: 새 알파 0건. 대신 '확정 알파를 실제로 무인 가동하는 운영 인프라'를 AWS에 실측 PASS로 깔았다 — 공식 1주 검증(2026-06-12~06-19)이 매시간 배치+텔레그램 알림+킬스위치 체제로 진행 중."),

("h1", "★Output of Chat (다음 채팅의 기준점)"),
("p", "① 공식 1주 가동 중: AWS C:\\run_Rauto에서 Rauto_Daily가 매시간 10분에 전체 재생 배치(test 재관통→check→daily_health→alert_check). 첫 라이브 페이퍼 거래 발생(2026-06-12 TS 1건, 일손익 -1.27%, bal_ts $11,212.09 — 워밍업 누적 $11,365.23에서). 주간 성적 기준선 = 6/12 0시 잔고 $11,365.23."),
("p", "② Stg16 OpsGuard 전축 AWS 실측 PASS: 텔레그램 알림 5종([PAPER] 태그·시작/진입/청산/오류·일일 하트비트), /status 4축 회신, /kill→kill.flag→태스크 비활성(자동복구 금지), 토큰은 OS 환경변수 전용(평문 grep 자체검사 영구 편입)."),
("p", "③ AWS 실측 버그 2건 해결: (a)alert_telegram URLError → v2(env 토큰 strip·form인코딩·무프록시 우선·reason 로그) (b)SYSTEM 태스크 python 못 찾음(0x80070002) → 풀경로 등록+run_daily.bat v7(RAUTO_PY env)."),
("p", "④ 다음 마일스톤: 2026-06-19 §5 주간 종합보고(신호동치·aux 커버 추이·주간 P&L/MDD vs 백테 동일구간·ER댐핑 발동기록) + 실계좌 소액 전환 체크리스트 본심사. 실계좌 전환의 절대 전제 = 상주 실시간 주문봇(체크리스트 B항) 별도 Stg 완성·검증."),
("p", "⑤ 확정 알파(§9) 변동 없음: 듀얼 = k0.77 + SW ER>=0.40×0.5 댐핑 → +1097.2% / MDD -16.24% / 표준6 CPCV 위반 0 (Stg12B). 이번 채팅은 이 알파를 '돌리는' 작업이었음."),

("h1", "0. 파일명·명명"),
("p", "본 보고서: Handover_07Prj_Ch4_RunAWS_stg16.docx (stg16 = 이번 채팅 최종 스테이지). 전 파일 영문명(zip 에러 방지). 약식 규칙: 07Prj_Ch4_... (Basic_Setup 4.5절)."),

("h1", "1. 작업 히스토리"),
("h2", "1.1 작업 파일 리스트 (전부 D:\\ml\\verify 하위)"),
("p", "[Stg14 폴더] run_daily.bat(v7: 단계 로그 append + RAUTO_PY 풀경로 폴백) / daily_health.py(v2: ★긴급 시 dauto_health.log에 ★RAUTO 미러) / AWS_DAILY_SETUP.txt(v2)"),
("p", "[Stg16 폴더 — 신설] ops_common.py / alert_telegram.py(v2) / alert_check.py(v2) / kill_guard.py / ops_status.py / telegram_poll.py / status_check.bat / test·check·run.bat / AWS_OPS_SETUP.txt(v2, Revoke 런북 포함) / aws_emergency_setup.bat / AWS_RDP_RUNBOOK_3MADI.txt / zip 5종(v1~v5 패치)"),
("p", "[기록] 00WorkHstr_INDEX.txt 약 15줄 추가(202606130042~202606131730) / TIL_Candidate_OpsRunbook_20260613.txt(Guide v4 후보, 승인 대기)"),
("h2", "1.2 코드 목적과 신뢰도 검증 방법"),
("p", "목적: 봇 본체(§8 해시 10종) 무수정 원칙 아래, 바깥에 운영 레이어(알림·킬·가시성)를 씌운다. 신뢰도 검증 3중: ①PC 샌드박스 8시나리오(mock HTTP·temp 경로·schtasks 미호출 — 실발송 0) ②오염검사 19항목(ops 해시 7종+봇 무수정 10종+토큰 평문 grep+결과 존재) ③AWS 실측(발사검증 200·/status 회신·schtasks 3종 Ready). 'PC PASS'는 참고일 뿐, AWS 실측만 PASS로 인정(캡틴 규칙)."),
("h2", "1.3 결과물 사용법 (캡틴 일상 운용)"),
("p", "아침 확인: 폰 텔레그램에 일일요약 1장(✅)이 와 있으면 정상. 거래 시 진입·청산 2장이 옴. 🚨가 오면 이상 — /status로 4축 확인. 긴급 정지: /kill 전송(해제는 수동: kill.flag+kill.flag.handled 삭제 후 schtasks /ENABLE). RDP에서는 status_check.bat 더블클릭."),
("h2", "1.4 비망록 — 캡틴이 강조·약속한 것"),
("p", "①설명은 축약 금지, 문장으로 풀어서(고딩 설명). 캡틴이 Claude Code 화면에 적응할 때까지 특히. — 메모리 박제됨 ②돈·시간에 닿는 설계 함의는 보고서에 묻지 말고 정면 질문으로 승인받기 ③AWS 실측만 PASS ④자동승인 모드라도 무기록 금지(변경 전 AUTO-CHANGE 1줄+INDEX 1줄) ⑤토큰 평문 어디에도 금지 ⑥실계좌 = 상주 주문봇 완성 후에만(재방송 배치로 실돈 금지)."),

("h1", "2. 시스템 아키텍처"),
("h2", "2.1 전체 그림 (AWS 1대, 윈도우 작업 스케줄러 4개)"),
("p", "Dauto_Collector(상주, 1분 수집) → C:\\BinanceData CSV → Rauto_Daily(매시간 10분: run_daily.bat가 test 재관통→check→daily_health→alert_check 순서로 실행, 멱등 전체재생이라 누락·갭 자동복원) → paper_ledger.csv·scorecard_daily.csv·stg14_health.log 갱신 → alert_check가 직전 상태(ops_state.json)와 diff해서 신규 이벤트만 텔레그램 발신. 별도로 Kill_Guard(1분)와 Telegram_Poll(1분)이 돈다."),
("h2", "2.2 모듈별 로직 (함수·변수는 Key노트에 코드 전체와 함께)"),
("p", "ops_common.py: 경로 해석의 단일 출처. rauto_dir()은 env RAUTO_DIR→C:\\run_Rauto→PC Stg14 상대경로→자기폴더 순서로 Stg14 산출물 폴더를 찾는다(PC·AWS 동일 코드 원칙). MODE_TAG는 RAUTO_MODE env가 LIVE면 [LIVE], 아니면 [PAPER]."),
("p", "alert_telegram.py v2: send(text) 하나. 토큰·chat_id는 환경변수에서만 읽고 strip으로 따옴표·공백 제거(①). 본문은 urlencode form(②), 무프록시 오프너 우선·실패 시 기본 경로(③), timeout 15(④). 미설정이면 NO-OP 로그만, 어떤 예외도 위로 던지지 않음(알림이 본 루틴을 못 깨게). 실패 로그에 reason 원문."),
("p", "alert_check.py v2: STEP4. ①시작(첫 OK 1회) ②③신규 원장행당 진입+청산 2장(원장은 완결거래 1행 구조라 실시간 진입 알림은 구조적으로 불가 — 상주 주문봇에서 해결 예정) ④오류(★긴급 신규·kill.flag) ⑤일일 하트비트(UTC 날짜 바뀐 첫 완주 1장). 중복 방지는 ops_state.json(seen 키=bot|entry_t, hb_date 등). 가격은 Dauto 1분봉 close 역참조."),
("p", "kill_guard.py: kill.flag 있으면 1회만 — 태스크 2종 /DISABLE + 청산 stub(페이퍼=로깅) + ★KILL을 health log와 텔레그램에. 마커(.handled)로 매분 재알림 방지. 자동복구 코드 자체가 없음(원칙)."),
("p", "telegram_poll.py: getUpdates 폴링. 등록 chat_id 외 전부 거부 로그. /status→ops_status.build_status() 4축 회신, /kill→flag 생성. 그 외 입력 무시(공격면 최소화). offset은 ops_poll_state.json."),
("p", "ops_status.py: 4축 = schtasks 2종 상태(로캘 무관 CSV 고정 컬럼순 파싱) + Dauto 최신행 시각 + health 마지막 줄 + scorecard 마지막 행."),
("h2", "2.3 주요 상수·환경변수"),
("p", "환경변수(전부 setx /M 필수): TELEGRAM_BOT_TOKEN / TELEGRAM_CHAT_ID(자격), RAUTO_DIR=C:\\run_Rauto(경로), RAUTO_PY(python 풀경로 — SYSTEM 태스크용), RAUTO_MODE(PAPER 기본/LIVE 전환 시). 테스트 전용: RAUTO_KILL_TEST, RAUTO_OPS_STATE 등(샌드박스 주입용). 상수: TASKS=[Rauto_Daily, Dauto_Collector](kill 대상), KILL_FLAG=C:\\BinanceData\\kill.flag."),

("h1", "3. 인수인계 시 주의점 (점프·미검증 명시)"),
("p", "[검증 필요 1 — 신뢰도 55] 매시간(HOURLY) 전환의 AWS 실측 미확인: schtasks 재등록 명령까지 전달했고 캡틴이 Y 응답 단계였음. 다음 채팅 첫 확인사항 — /status에서 Rauto_Daily Last가 직전 시각대 10분·rc 0인지."),
("p", "[검증 필요 2 — 신뢰도 55] 일일 하트비트의 AWS 실발신 미확인(PC mock만 PASS). 패치 zip(Stg16_patch_v5_hourly_heartbeat.zip) 적용 여부도 함께 확인할 것."),
("p", "[검증 필요 3 — 신뢰도 55] scorecard 2026-06-12 행의 oi_blunt_pct=73.67(이전 100). 의미(뭉툭화 비율 하락=라이브 OI 정상 유입 시작) 해석은 합리적이나 수치 추이를 6-19 보고에서 검증할 것."),
("p", "[구조적 한계 — 확정] 페이퍼 단계 알림은 '봉마감 후 다음 배치'에서 발신(매시간 체제로 최대 1시간 지연). 거래 판정 자체는 봉마감 봇이라 상주봇과 1원까지 동일 — 이건 버그가 아니라 설계이며, 실계좌는 상주 주문봇이 전제."),
("p", "★알파 저장 규칙 리마인드(지침): 알파가 확인되면 반드시 정리해서 G:\\내 드라이브\\00AI개발지식DB\\자산관리\\유동자산\\자동매매\\06 ChampBot\\00ALPHA_Confirm_Bot 폴더에 저장하라고 캡틴에게 강조할 것. (이번 채팅은 새 알파 0건 — 해당 없음. §9 듀얼 확정 알파는 기존 NOTE_OfficialWeek_Start_20260612.txt로 G드라이브 저장 완료 상태.)"),

("h1", "4. 개발 토론·문제해결 과정 (상세·코드 전체는 Key노트)"),
("p", "①머신 불일치 3회: AWS 지시가 PC 세션에 반복 수신 → hostname·Test-Path 증거로 매번 중단 보고 → 최종 'RDP 복붙 런북(3마디·정지점)' 방식으로 합의. ②발신불능 URLError: 판별표(가설A 토큰영역/가설B PATH) 1왕복으로 원인 좁힘 → v2 패치 적용 → 실측 200. ③SYSTEM python: Last Result 0x80070002로 B 확정 → 풀경로 재등록+bat v7. ④'재방송' 논쟁: 일배치=알림 지연 함의를 정면 질문하지 않은 과실 인정 → 매시간+하트비트로 수습, 실계좌=상주봇 전제 명문화. ⑤보안: 토큰 1차 채팅 노출 → BotFather Revoke 런북 신설+평문 grep을 오염검사 영구 항목화."),

("h1", "5. 다음 채팅 작업 계획 (Output of Chat 기준)"),
("p", "1순위: 06-19 §5 주간 종합보고 — 기준선 $11,365.23(6/12 0시), 항목=신호동치 7/7·★긴급 0·aux 커버 추이(oi_z·atr)·주간 P&L/MDD vs 백테 동일구간·ER댐핑/필터 발동기록·(참고)$10k 정규화 환산. 2순위: 실계좌 소액 전환 체크리스트 본심사(게이트 4조건). 3순위: 상주 실시간 주문봇 설계 Stg 발주(체크리스트 B항 6~10 구현 — 멱등 발주·포지션 대사·하드스톱). 보류 트리거: 듀얼 k 상향 재평가(라이브 3개월), rclone/S3(라이브 시)."),

("h1", "6. 다음 채팅 시작 시 비판적 검토 사항"),
("p", "①HOURLY 전환·하트비트 실측 확인(주의점 1·2) ②매시간 체제 부작용 점검: stg14_health.log 하루 24줄 증가·AWS CPU/디스크 여유(캡틴 言 'AWS 느리고 용량 적음') ③ops_state.json 유실 시 알림 전량 재발신 위험 — 백업 또는 재생성 가드 검토 ④Dauto_Collector가 kill 대상에 포함된 게 적절한가(데이터 수집은 살려둘지 — 캡틴 결정) ⑤slip_p50 컬럼이 빈 값(페이퍼라 체결 슬리피지 없음) — 6-19 보고의 '슬리피지 vs 1.7bp' 항목은 측정 불가임을 미리 인지."),

("h1", "부록 A. Key 파일 리스트"),
("p", "[AWS C:\\run_Rauto 운영] run_daily.bat(v7) / ops 7종 py / status_check.bat / bots\\(봇 본체 10종 — §8 무수정) [PC 원본] D:\\ml\\verify\\07Prj_Ch4_RunAWS_Stg16_OpsGuard\\ 전부 [런북] AWS_RDP_RUNBOOK_3MADI.txt·AWS_OPS_SETUP.txt(v2 Revoke 런북) [코드 전체] KeyNote_07Prj_Ch4_RunAWS_stg16_OpsGuard.docx"),

("h1", "부록 B. 신뢰도 점수표 (출처 명시)"),
("p", "95: AWS 발사검증 200·/status 4축 회신·schtasks 3종 Ready(캡틴 RDP 실측 출력) / 봇 본체 무수정(해시 10종 전수일치) / 평문 토큰 0건(grep 전수) / 첫 라이브 거래 1건·잔고 추이(scorecard 실측)"),
("p", "55: HOURLY 전환 완료 여부 / 하트비트 AWS 실발신 / oi_blunt 73.67 해석 — 전부 다음 실측에서 확정"),
("p", "15: URLError의 정확한 단일 원인(①strip인지 ③프록시인지 — v2가 동시 차단해 개별 특정은 안 됨. 운영상 무해하므로 추적 불요)"),
]


def build_handover():
    doc = docx.Document()
    for style, text in HANDOVER:
        add(doc, style, text)
    out = os.path.join(PKG, "Handover_07Prj_Ch4_RunAWS_stg16.docx")
    doc.save(out)
    return out


KEY_FILES = [
    (S16, "ops_common.py"), (S16, "alert_telegram.py"), (S16, "alert_check.py"),
    (S16, "kill_guard.py"), (S16, "ops_status.py"), (S16, "telegram_poll.py"),
    (S16, "status_check.bat"), (S16, "check_07Prj_Ch4_RunAWS_Stg16_OpsGuard.py"),
    (S16, "test_07Prj_Ch4_RunAWS_Stg16_OpsGuard.py"),
    (S14, "run_daily.bat"), (S14, "daily_health.py"),
]


def build_keynote():
    doc = docx.Document()
    add(doc, "h1", "KeyNote_07Prj_Ch4_RunAWS_stg16_OpsGuard — 핵심·검증 로직 + 코드 전체")
    add(doc, "p", "작성 2026-06-13. 각 파일 머리 주석에 목적·버전·변경사유가 있으니 코드와 함께 읽을 것. "
                  "검증 로직 핵심: test 8시나리오는 전부 샌드박스(temp 경로·mock HTTP·schtasks 미호출)라 "
                  "PC에서 안전하게 반복 가능. 실발송·실등록은 AWS에서만, 캡틴 승인 단위로.")
    add(doc, "h2", "고딩 설명 — 이 레이어가 하는 일")
    add(doc, "p", "봇 본체는 손대지 않는다. 대신 봇이 남기는 파일(원장·스코어카드·헬스로그)을 '읽기만' 해서 "
                  "변화가 있으면 폰으로 알려주고(알림), 폰에서 명령이 오면 파일 하나(kill.flag)를 만들어 "
                  "봇을 멈춘다(킬). 봇과 운영 레이어가 파일로만 대화하므로, 운영 레이어가 어떤 버그를 내도 "
                  "봇 계산은 오염되지 않는다 — 이것이 '엔진 무수정' 원칙의 실전 형태다.")
    for d, fn in KEY_FILES:
        p = os.path.join(d, fn)
        add(doc, "h2", f"[코드 전체] {fn}  ({'Stg16' if d == S16 else 'Stg14'})")
        try:
            body = open(p, encoding="utf-8").read()
        except OSError as e:
            body = f"(읽기 실패: {e})"
        add(doc, "code", body)
    out = os.path.join(PKG, "KeyNote_07Prj_Ch4_RunAWS_stg16_OpsGuard.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    print("handover:", build_handover())
    print("keynote:", build_keynote())

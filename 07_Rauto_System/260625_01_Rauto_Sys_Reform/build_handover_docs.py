# -*- coding: utf-8 -*-
# [build_handover_docs.py] 세션 260625_01_Rauto_Sys_Reform 인수인계 문서 생성 (§17 A).
#   Handover·KeyNote·Work_Order 신규 + Hstr_Ver_Up ADR 추가 + Guide v6 TIL 추가.
import os
from docx import Document
from docx.shared import Pt, RGBColor
DOCS = r"D:\ML\RfRauto\07_Rauto_System\260625_01_Rauto_Sys_Reform"
NOTE = r"D:\ML\RfRauto\01_제작노트"
BSP = r"D:\ML\RfRauto\00_Basic_Setup_Package"
HV = os.path.join(NOTE, "0004Hstr_Ver_Up_Bots", "Hstr_Ver_Up_Rauto_Champion_Sys.docx")
GUIDE = os.path.join(BSP, "Guide_AlphaDiscovery_Method_TrendStack_Upgrade_v6.docx")


def H(doc, t, lv=1):
    h = doc.add_heading(t, level=lv); return h

def P(doc, t, b=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = b; r.font.size = Pt(10.5); return p

def tbl(doc, rows, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, c in enumerate(row):
            cell = t.cell(i, j); cell.text = str(c)
            for para in cell.paragraphs:
                for run in para.runs: run.font.size = Pt(9); run.bold = (i == 0 and header)
    return t


# ───────── 1. Handover ─────────
d = Document()
d.add_heading("Handover — 260625_01_Rauto_Sys_Reform", 0)
P(d, "세션ID: 260625_01_Rauto_Sys_Reform · 날짜: 2026-06-25 · 유형: 시스템 아키텍처 전환점(§14)", True)

H(d, "1. 세션 개요")
P(d, "직전 세션(D:\\ML\\verify 폴더, RevoiExitRegime)을 RfRauto 폴더로 더블클릭 전환하며 세션이 안 보인 것을 복구로 시작. "
     "REVoi 청산 1폴드 닫기·실비용 재산정을 거치며 '비용 갑/을·슬립 재논쟁'이 반복되자, 캡틴이 근본해법으로 "
     "Rauto를 5모듈로 나누는 구조개혁을 지시. 토론·위험검증·승인 후 ①②③ 모듈을 추출·앵커검증까지 완료.")

H(d, "2. Output of Chat (핵심 산출)")
P(d, "★Rauto 5모듈 구조개혁 승인 + 핵심 3모듈 추출·검증 완료(전환점).", True)
tbl(d, [
    ["모듈", "내용", "상태"],
    ["[0] 관제센터", "중앙 1m·룩어헤드 게이트·슬롯/챔피언", "②데이터층 완료, 슬롯/챔피언=다음"],
    ["[1] 매매신호", "rauto_signal.py (rev_side 래퍼, Signal만)", "완료"],
    ["[2] 매매결정", "rauto_decision.py (좁은범위 '가': 사이징·SL)", "완료"],
    ["[3] RautoCEX", "rauto_cex.py (Fee/Slip/Fill/Margin, 비용 단일출처)", "완료·앵커PASS"],
    ["[4] 결과분석", "Back2TV (기존)", "기존 유지"],
])
P(d, "★검증: 매 단계 앵커(+1851.6%/MDD-24.6%)가 1원단위 동일. 비용 4곳 산재→RautoCEX 한곳. 룩어헤드 게이트로 미래참조 0건(게이트 없으면 100% 누수).")

H(d, "3. 작업 경과")
for s in ["세션 복구: 직전 작업이 D:\\ML\\verify 폴더(RevoiExitRegime, 13:49)에 있었음을 확인, RfRauto에서 파일기반으로 이어감.",
          "REVoi 1폴드 닫기: 레짐스톱(TIGHT) 1.4→1.6이 표준6 위반 7%→0% 닫는 진짜 레버 발견(FoldClose).",
          "실비용 재산정: +1852%는 슬립0 낙관. 내가 지어낸 슬립모델(4.7bp)은 과대→폐기. 캡틴 측정 갭슬립=0bp, 스프레드1bp만 더하면 현실 +1483%(SlipRecheck).",
          "구조개혁: 비용/룩어헤드 재논쟁의 근본원인=관심사 미분리. 5모듈안 제안→선행연구(이벤트드리븐·NautilusTrader) 대조→위험 8시나리오 검증→7안전장치 조건부 승인.",
          "①②③ 추출: RautoCEX·DataHub·Signal/Decision 분리, 각 단계 앵커 회귀 PASS."]:
    P(d, "· " + s)

H(d, "4. 핵심 기술·코드")
P(d, "rauto_cex.py: bt_full이 R에 박은 기본비용(MK+TK+펀딩)을 gross로 복원→자기 FeeModel/SlipModel/MarginModel로 재차감 = 비용 단일출처. 슬립0이면 기존과 동일, 스프1bp면 +1483%.")
P(d, "rauto_datahub.py: resample는 label='left'(라벨=봉시작). 게이트=봉마감(라벨+TF분)<=now인 봉만 공개. 게이트 없는 라벨접근은 100% 미래참조 누수(실측).")
P(d, "rauto_signal/decision.py: 검증엔진(rev_side·gen_trades) 래퍼만, 무수정(§8)·재구현 없음(§15.1). 결정모듈 범위=캡틴 (가)좁게 채택.")

H(d, "5. 신뢰도표")
tbl(d, [
    ["항목", "수치/판정", "신뢰도"],
    ["①RautoCEX 무손상추출", "기존≡신모듈 +1851.6% 차이 0.000%p", "95"],
    ["②룩어헤드 게이트", "미래참조 누수 0/1000 (게이트없으면 100%)", "95"],
    ["③신호/결정 분리 체인", "전체배선 +1851.6% 차이 0.000%p", "95"],
    ["REVoi 실비용 재산정", "+1852%→+1483%(스프1bp, 갭슬립0 측정)", "85"],
    ["FoldClose 1폴드", "TIGHT1.6 위반7%→0%(held-out)", "85"],
])

H(d, "6. 산출물 위치")
P(d, "단일출처 SPEC: 07_Rauto_System\\260625_01_Rauto_Sys_Reform\\260625_01_Rauto_Sys_Reform_SPEC.md")
P(d, "공용엔진 승급: 04_공용엔진코드\\engines\\rauto_cex.py · rauto_datahub.py")
P(d, "v0 모듈·검증·그래프·zip: 07_Rauto_System\\260625_01_Rauto_Sys_Reform\\ (signal/decision·AnchorTest·LookAheadTest·WiredAnchorTest·ModuleMap·RiskMatrix)")
P(d, "REVoi 백테 산출: 00_WorkHstr\\BackTest_Output\\260625_09~12_*")

H(d, "7. 다음 세션 1수")
P(d, "④ 관제센터(슬롯·챔피언) 설계 — 챔피언선발=CPCV/held-out 게이트(안전장치5), full표본 금지. + RautoCEX Live mode·tick 슬립보정(안전장치7)·signal/decision engines 승급(deps정리). 상세=Work_Order_RautoSysReform_20260625.")
d.save(os.path.join(DOCS, "260625_01_Rauto_Sys_Reform_Handover.docx"))
print("[저장] Handover")

# ───────── 2. KeyNote ─────────
d = Document()
d.add_heading("KeyNote — 260625_01_Rauto_Sys_Reform (기술 상세)", 0)
H(d, "A. 비용 2레이어 철칙 (§7)")
P(d, "selection_cost(신호 4bp, 봉선정용·P&L아님)와 execution_cost(RautoCEX, 진짜P&L)는 다른 것. RautoCEX로 모으는 건 execution만. 이걸 어기면 갑/을 부활. 코드에서 이름 분리.")
H(d, "B. 앵커 회귀 규율 (안전장치1)")
P(d, "모든 모듈 추출은 '같은 config→같은 수익(+1851.6%/MDD-24.6%)'을 1원단위로 재현해야 머지. ①②③ 모두 차이 0.000%p로 통과. 이게 리팩터 회귀(S3)·집중오염(S1)의 안전망.")
H(d, "C. 룩어헤드 100% 누수 발견 (안전장치3)")
P(d, "resample label='left'라 봉 라벨=시작시각. 게이트 없이 '라벨<=now'로 접근하면 그 봉은 마감 전까지 항상 미래로 뻗어 100% 미래참조. DataHub의 'close_time<=now' 게이트가 0건으로 차단.")
H(d, "D. 실비용 재산정 (캡틴 교정)")
P(d, "+1852%는 청산 슬립0 낙관. AI가 지어낸 슬립모델(base3+분위, 청산4.7bp)은 과대→폐기. 캡틴 측정(exec_realism 932거래) 갭슬립=0.00bp. 1m봉이 못 보는 호가스프레드 1bp만 더하면 현실 +1483%(MDD-25 유지). 진입·구조청산은 반전예상 정해진레벨=지정가 메이커(체결가능 1m사전판정).")
H(d, "E. 7 안전장치 (위험 S1~S8 대책)")
for s in ["①앵커회귀관문 ②FillModel철칙(도달≠체결,스톱=시장가+슬립) ③봉마감게이트+룩어헤드단위테스트 ④비용 이름분리",
          "⑤챔피언선발=CPCV/held-out만(full표본금지) ⑥벡터·이벤트 모델공유 ⑦슬립 틱실측보정"]:
    P(d, "· " + s)
d.save(os.path.join(DOCS, "260625_01_Rauto_Sys_Reform_KeyNote.docx"))
print("[저장] KeyNote")

# ───────── 3. Work_Order ─────────
d = Document()
d.add_heading("Work Order — Rauto 구조개혁 후속 (2026-06-25)", 0)
P(d, "맥락: 5모듈 구조개혁 ①②③ 완료(전환점). ④와 보강 과제를 다음 세션 이후로.", True)
tbl(d, [
    ["과제", "내용", "완료조건", "재평가"],
    ["④ 관제센터", "슬롯/챔피언 매니저 + DataHub 연결 끝단배선", "REVoi 1봇 끝단 데모 + 챔피언선발 CPCV게이트", "다봇 생긴 후"],
    ["RautoCEX Live", "Sim↔Live 인터페이스의 Live 구현(실거래소 체결수신)", "테스트넷 체결과 Sim 비용모델 오차 측정", "테스트넷 단계"],
    ["슬립 틱보정(안전장치7)", "aggTrades 실측으로 SlipModel 계수 보정", "스프레드/충격 실측 대체, +1483% 갱신", "틱데이터 확보 후"],
    ["engines 승급", "signal/decision의 research deps(rev_side·gen_trades) 정리 후 engines 이전", "deps 끊고 engines import만으로 동작", "다음 세션"],
    ["연구경로 적용", "벡터 백테(연구)도 RautoCEX 모델 공유(안전장치6)", "벡터·이벤트 1회 교차검증 일치", "④ 후"],
])
d.save(os.path.join(DOCS, "Work_Order_RautoSysReform_20260625.docx"))
print("[저장] Work_Order")

# ───────── 4. Hstr_Ver_Up ADR 추가 ─────────
try:
    d = Document(HV)
    d.add_heading("ADR — Rauto 5모듈 구조개혁 (2026-06-25, 세션 260625_01_Rauto_Sys_Reform)", 1)
    for line in [
        "①문제제기: 비용(14/8/4bp)·슬립 재논쟁이 봇 수정 때마다 반복 → 수치가 매번 바뀌고 캡틴 신뢰 훼손.",
        "②원인분석: 관심사 미분리 — 비용/체결 로직이 4곳(신호·실행·bt_full·fib_replay)에 산재, 데이터·신호·결정·체결이 한 루프에 혼재.",
        "③대안검토: (a)그때그때 비용확인=현행, 기각 (b)★5모듈 분리+RautoCEX 단일비용=채택. 선행연구(이벤트드리븐·NautilusTrader Fill/Slip/Fee·backtest-live parity)와 일치.",
        "④해결법: [0]관제센터 [1]신호 [2]결정 [3]RautoCEX(체결+비용) [4]분석. 7안전장치(앵커회귀·FillModel철칙·봉마감게이트·비용이름분리·챔피언CPCV게이트·모델공유·틱보정). 착수=RautoCEX부터 앵커검증.",
        "⑤결론: ①②③ 추출 완료, 앵커 +1851.6% 1원단위 재현. 비용 단일출처·룩어헤드 0건 달성. 단일출처=260625_01_Rauto_Sys_Reform_SPEC.md.",
        "⑥향후적용: ④관제센터·Live·틱보정은 Work_Order. 모든 신규 봇은 이 모듈경계로 붙인다(신호=Signal만·비용=RautoCEX만).",
    ]:
        P(d, line)
    d.save(HV)
    print("[추가] Hstr_Ver_Up ADR")
except Exception as e:
    print("[건너뜀] Hstr_Ver_Up:", e)

# ───────── 5. Guide v6 TIL 추가 ─────────
try:
    d = Document(GUIDE)
    d.add_heading("TIL (2026-06-25, 세션 260625_01_Rauto_Sys_Reform) — 백테 비용·검증 방법론", 1)
    for line in [
        "TIL-1 [비용]: 화려한 백테수익은 대개 '슬립0 낙관'. +1852%(REVoi MDD25)는 청산 시장가슬립을 0으로 본 값. 출처·신뢰85.",
        "TIL-2 [측정≠충분]: 1m 갭슬립 측정(exec_realism)은 갭관통=0으로 정확하나, 호가스프레드(sub-1분)는 구조적으로 못 봄. 필요(갭)+스프레드 바닥(1bp)이라야 충분. 갭0+스프1bp→현실 +1483%.",
        "TIL-3 [지정가 사전판정]: 반전예상 '정해진 레벨' 진입/구조청산은 1m 도달여부로 체결가능을 미리 알 수 있음=지정가 메이커. 시장가는 스톱(fibstop)만.",
        "TIL-4 [모듈화는 일관성이지 정답이 아님]: 관심사 분리(RautoCEX)는 '비용 재논쟁'을 끝내지만 '진짜 슬립이 얼마냐'는 별개로 실측보정 필요. 신뢰90.",
        "TIL-5 [앵커 회귀]: 어떤 리팩터도 '같은 config→같은 수익 1원단위 재현'을 통과해야 무손상. 통과 못하면 멈추고 원인규명(§15.2).",
    ]:
        P(d, line)
    d.save(GUIDE)
    print("[추가] Guide v6 TIL")
except Exception as e:
    print("[건너뜀] Guide:", e)

print("[완료] 인수인계 문서 생성")
